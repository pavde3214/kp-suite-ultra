# -*- coding: utf-8 -*-
# integrate_docs_and_full_kp_page.py
from pathlib import Path
import sqlite3, re, textwrap

BASE = Path(__file__).resolve().parent

def w(rel, content):
    p = BASE / rel
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(textwrap.dedent(content).lstrip("\n"), encoding="utf-8")
    print("wrote:", rel)

def patch_file(path, repl_func):
    p = BASE / path
    s = p.read_text(encoding="utf-8")
    ns = repl_func(s)
    if ns != s:
        p.write_text(ns, encoding="utf-8")
        print("patched:", path)
    else:
        print("no changes:", path)

# 0) Гарантируем каталоги
for d in ["storage/templates/contracts", "storage/docs"]:
    (BASE / d).mkdir(parents=True, exist_ok=True)

# 1) Миграция клиентов: добавим реквизиты
db_path = BASE / "kp.db"
if db_path.exists():
    con = sqlite3.connect(db_path)
    cur = con.cursor()
    cols = {r[1] for r in cur.execute("PRAGMA table_info(customers)").fetchall()}
    need = [
        ("company_name","TEXT"),("inn","TEXT"),("kpp","TEXT"),("ogrn","TEXT"),
        ("bank_name","TEXT"),("bank_bik","TEXT"),("bank_acc","TEXT"),("bank_corr","TEXT"),
        ("position","TEXT"),("contact_person","TEXT")
    ]
    for col, typ in need:
        if col not in cols:
            cur.execute(f"ALTER TABLE customers ADD COLUMN {col} {typ}")
            print("ALTER customers ADD", col)
    con.commit(); con.close()

# 2) Перепишем модель Customer с новыми полями
w("app/models/customer.py", """
from sqlalchemy import Column, Integer, String
from app.db.base import Base
class Customer(Base):
    __tablename__ = "customers"
    id = Column(Integer, primary_key=True)
    fullname = Column(String(200), nullable=False)
    passport = Column(String(200))
    address = Column(String(300))
    address_object = Column(String(300))
    phone = Column(String(100))
    email = Column(String(150))
    # реквизиты
    company_name = Column(String(200))
    inn = Column(String(50))
    kpp = Column(String(50))
    ogrn = Column(String(50))
    bank_name = Column(String(200))
    bank_bik = Column(String(50))
    bank_acc = Column(String(50))
    bank_corr = Column(String(50))
    position = Column(String(100))
    contact_person = Column(String(150))
""")

# 3) Обновим карточку клиента (полная)
w("app/web/templates/customer_card_full.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Карточка клиента #{{ c.id }}</h2>
<div class="grid">
  <form class="card" method="post" action="/customers/{{ c.id }}/save">
    <h3>Основные данные</h3>
    <div class="grid2">
      <label>ФИО/Наименование:</label><input name="fullname" value="{{ c.fullname or '' }}">
      <label>Паспорт:</label><input name="passport" value="{{ c.passport or '' }}">
      <label>Телефон:</label><input name="phone" value="{{ c.phone or '' }}">
      <label>Email:</label><input name="email" value="{{ c.email or '' }}">
      <label>Адрес регистрации:</label><input name="address" value="{{ c.address or '' }}">
      <label>Адрес объекта:</label><input name="address_object" value="{{ c.address_object or '' }}">
    </div>

    <h3>Реквизиты (для договора)</h3>
    <div class="grid2">
      <label>Компания (если юр. лицо):</label><input name="company_name" value="{{ c.company_name or '' }}">
      <label>Должность подписанта:</label><input name="position" value="{{ c.position or '' }}">
      <label>Контактное лицо:</label><input name="contact_person" value="{{ c.contact_person or '' }}">
      <label>ИНН / КПП:</label>
      <div style="display:flex; gap:8px">
        <input name="inn" value="{{ c.inn or '' }}" placeholder="ИНН" style="width:200px">
        <input name="kpp" value="{{ c.kpp or '' }}" placeholder="КПП" style="width:200px">
      </div>
      <label>ОГРН/ОГРНИП:</label><input name="ogrn" value="{{ c.ogrn or '' }}">
      <label>Банк:</label><input name="bank_name" value="{{ c.bank_name or '' }}">
      <label>БИК / р/с / к/с:</label>
      <div style="display:flex; gap:8px">
        <input name="bank_bik" value="{{ c.bank_bik or '' }}" placeholder="БИК" style="width:180px">
        <input name="bank_acc" value="{{ c.bank_acc or '' }}" placeholder="р/с" style="width:280px">
        <input name="bank_corr" value="{{ c.bank_corr or '' }}" placeholder="к/с" style="width:280px">
      </div>
    </div>

    <div style="margin-top:8px">
      <button>Сохранить</button>
      <a class="btn" href="/proposals/create?customer_id={{ c.id }}">Создать КП</a>
      <a class="btn" href="/proposals/new?customer_id={{ c.id }}">Быстрое КП</a>
      <a class="btn" href="/contract/editor?customer_id={{ c.id }}{% if latest_kp %}&kp_id={{ latest_kp.id }}{% endif %}">Редактор договора</a>
      <a class="btn" href="/contracts/create-all?customer_id={{ c.id }}{% if latest_kp %}&kp_id={{ latest_kp.id }}{% endif %}" target="_blank">Сформировать пакет (Договор + Прил.1–6)</a>
    </div>
  </form>

  <div class="card">
    <h3>КП клиента</h3>
    <ul>
      {% for p in proposals %}
        <li><a href="/proposals/{{ p.id }}/edit">{{ p.created_at.strftime('%Y-%m-%d') if p.created_at else '' }} — {{ p.number }} — {{ p.title }}</a></li>
      {% else %}
        <li>КП пока нет</li>
      {% endfor %}
    </ul>
  </div>
</div>
{% endblock %}
""")

# 4) Обновим сохранение клиента (main.py)
def patch_main_save_customer(s):
    s = re.sub(
        r"@app.post\(\"/customers/\{cid\}/save\"\)[\s\S]*?return RedirectResponse[^\n]+\n",
        textwrap.dedent("""
        @app.post("/customers/{cid}/save")
        async def customers_save(cid: int, request: Request):
            form = await request.form()
            db = SessionLocal()
            c = db.get(Customer, cid)
            if not c:
                db.close()
                return RedirectResponse(url="/customers", status_code=303)
            # базовые
            c.fullname = form.get("fullname") or c.fullname
            c.passport = form.get("passport") or c.passport
            c.address = form.get("address") or c.address
            c.address_object = form.get("address_object") or c.address_object
            c.phone = form.get("phone") or c.phone
            c.email = form.get("email") or c.email
            # реквизиты
            c.company_name = form.get("company_name") or c.company_name
            c.position = form.get("position") or c.position
            c.contact_person = form.get("contact_person") or c.contact_person
            c.inn = form.get("inn") or c.inn
            c.kpp = form.get("kpp") or c.kpp
            c.ogrn = form.get("ogrn") or c.ogrn
            c.bank_name = form.get("bank_name") or c.bank_name
            c.bank_bik = form.get("bank_bik") or c.bank_bik
            c.bank_acc = form.get("bank_acc") or c.bank_acc
            c.bank_corr = form.get("bank_corr") or c.bank_corr
            db.add(c); db.commit(); db.close()
            return RedirectResponse(url=f"/customers/card/{cid}", status_code=303)
        """).lstrip("\n"),
        flags=re.M
    )
    return s

patch_file("app/main.py", patch_main_save_customer)

# 5) Улучшенная страница «Создать КП»: список клиентов + форма нового
w("app/web/templates/proposal_new.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Создание КП</h2>
<div class="card">
  <form method="post" action="/proposals/new">
    <h3>Клиент</h3>
    <div class="grid2">
      <label>Выбрать из списка:</label>
      <select name="customer_id">
        <option value="">— создать нового —</option>
        {% for c in customers %}
          <option value="{{ c.id }}" {% if request.query_params.get('customer_id') and request.query_params.get('customer_id')|int == c.id %}selected{% endif %}>
            [{{ c.id }}] {{ c.fullname }} — {{ c.phone or '' }}
          </option>
        {% endfor %}
      </select>
      <label>ФИО/Наименование (для нового):</label><input name="fullname" placeholder="ООО «Ромашка» / Иванов Иван" value="">
      <label>Телефон / Email:</label>
      <div style="display:flex; gap:8px">
        <input name="phone" placeholder="+7 ..." style="width:220px">
        <input name="email" placeholder="mail@example.com" style="width:280px">
      </div>
      <label>Адрес регистрации:</label><input name="address" placeholder="Город, улица ...">
      <label>Адрес объекта:</label><input name="address_object" placeholder="Адрес, который попадёт в КП/Договор">
    </div>

    <h3>Параметры КП</h3>
    <div class="grid2">
      <label>Название КП:</label><input name="title" value="Коммерческое предложение">
      <label>Адрес объекта (если переопределить):</label><input name="site_address">
      <label>Шапка (HTML):</label><textarea name="header_html" rows="3"><b>VENT•PRO HVAC</b></textarea>
    </div>
    <button>Создать КП</button>
  </form>
</div>
{% endblock %}
""")

# 6) Поменяем обработчик GET /proposals/new, чтобы отдавал список клиентов
def patch_main_proposal_new(s):
    s = re.sub(
        r"@app.get\(\"/proposals/new\"[^\n]+\)\s*def new_proposal_page[^\n]*\:[\s\S]*?return templates\.TemplateResponse\([^\)]+\)",
        textwrap.dedent("""
        @app.get("/proposals/new", response_class=HTMLResponse)
        def new_proposal_page(request: Request, customer_id: int | None = None):
            db = SessionLocal()
            customers = db.execute(select(Customer).order_by(Customer.id.desc())).scalars().all()
            db.close()
            return templates.TemplateResponse("proposal_new.html", {"request": request, "customers": customers})
        """).lstrip("\n"),
        flags=re.M
    )
    return s

patch_file("app/main.py", patch_main_proposal_new)

# 7) Роут «сформировать полный пакет» (UI) + API
# 7.1) сервис: аккуратная подстановка плейсхолдеров если они есть, иначе просто копия
w("app/services/appendices.py", """
from pathlib import Path
from docx import Document
from app.services.docx_merge import render_docx

def render_or_copy(src: Path, dst: Path, mapping: dict):
    try:
        # если в шаблоне есть [[...]] — рендерим
        doc = Document(str(src))
        has_ph = False
        for p in doc.paragraphs:
            if "[[" in p.text and "]]" in p.text:
                has_ph = True; break
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    if "[[" in c.text and "]]" in c.text:
                        has_ph = True; break
        if has_ph:
            render_docx(str(src), str(dst), mapping)
        else:
            dst.write_bytes(src.read_bytes())
    except Exception:
        # на любой сбой — просто копия
        dst.write_bytes(src.read_bytes())
    return str(dst)
""")

# 7.2) API: /api/v1/contracts/create-all
w("app/api/routers/contracts.py", """
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import select
from pathlib import Path
import json
from datetime import date
from app.api.deps import get_db
from app.models import Customer, Owner, Estimate, Proposal, DocumentRec, ContractData
from app.services.numbering import generate_contract_number
from app.services.payments import estimate_totals, payment_schedule_100_70_30, proposal_totals
from app.services.docx_merge import render_docx
from app.services.spec_builder import build_spec_docx
from app.services.export_xlsx import export_spec_xlsx
from app.services.appendices import render_or_copy
from app.config import TEMPLATES_DIR, DOCS_DIR

router = APIRouter(prefix="/api/v1/contracts", tags=["contracts"])

def _fmt(x):
    try: return f"{float(x):,.0f}".replace(",", " ")
    except: return str(x)

def build_ctx(db: Session, customer_id: int, estimate_id: int | None, kp_id: int | None, overrides: dict | None = None):
    cust = db.get(Customer, customer_id)
    if not cust: raise HTTPException(404, "customer not found")
    owner = db.get(Owner, 1)
    if not owner:
        owner = Owner(id=1, display_name="Исполнитель"); db.add(owner); db.commit(); db.refresh(owner)
    est = db.get(Estimate, estimate_id) if estimate_id else None
    kp = db.get(Proposal, kp_id) if kp_id else None
    totals = {"equipment":0,"labor":0,"total":0}; sched = []
    if est:
        totals = estimate_totals(est); sched = payment_schedule_100_70_30(est)
    elif kp:
        totals = proposal_totals(kp); sched = [{"title":"Оплата по КП","amount":totals["total"]}]
    number = generate_contract_number(db)
    ctx = {
        "CONTRACT_NO": number, "CONTRACT_DATE": date.today().strftime("%d.%m.%Y"),
        "CUSTOMER_FIO": cust.fullname or "", "CUSTOMER_PASSPORT": cust.passport or "",
        "CUSTOMER_ADDR_REG": cust.address or "", "CUSTOMER_PHONE": cust.phone or "",
        "CUSTOMER_EMAIL": cust.email or "",
        "CUSTOMER_COMPANY": cust.company_name or "",
        "CUSTOMER_POSITION": cust.position or "",
        "CUSTOMER_CONTACT": cust.contact_person or "",
        "CUSTOMER_INN": cust.inn or "", "CUSTOMER_KPP": cust.kpp or "", "CUSTOMER_OGRN": cust.ogrn or "",
        "CUSTOMER_BANK": cust.bank_name or "", "CUSTOMER_BIK": cust.bank_bik or "",
        "CUSTOMER_ACC": cust.bank_acc or "", "CUSTOMER_CORR": cust.bank_corr or "",
        "OBJECT_ADDR": (getattr(est, "site_address", None) if est else None) or (kp.site_address if kp else None) or (cust.address_object or ""),
        "EXECUTOR_FIO": owner.fullname or owner.display_name or "Исполнитель",
        "EXECUTOR_PASSPORT": owner.passport or "", "EXECUTOR_ADDR": owner.address or "",
        "EXECUTOR_PHONE": owner.phone or "", "EXECUTOR_EMAIL": owner.email or "",
        "EXECUTOR_INN": owner.inn or "", "EXECUTOR_OGRN": owner.ogrn or "",
        "SUM_EQUIP": _fmt(totals["equipment"]), "SUM_WORK":  _fmt(totals["labor"]),
        "SUM_TOTAL": _fmt(totals["total"]),
        "STAGE1": _fmt(sched[0]["amount"] if sched else 0),
        "STAGE2": _fmt(sched[1]["amount"] if len(sched)>1 else 0),
        "STAGE3": _fmt(sched[2]["amount"] if len(sched)>2 else 0),
    }
    if overrides:
        ctx.update({k: v for k, v in overrides.items()})
    return number, ctx, est, kp

@router.post("/create")
def create_contract(customer_id: int = Query(...), estimate_id: int | None = Query(None), kp_id: int | None = Query(None), db: Session = Depends(get_db)):
    number, ctx, est, kp = build_ctx(db, customer_id, estimate_id, kp_id)
    dog_tpl = TEMPLATES_DIR / "Договор.docx"
    if not dog_tpl.is_file(): raise HTTPException(400, "Нет шаблона Договор.docx в storage/templates/contracts")
    contract_out = DOCS_DIR / f"Договор_{number}.docx"
    render_docx(str(dog_tpl), str(contract_out), ctx)
    app1_out = None
    if est:
        app1_out = Path(build_spec_docx(est, number)); export_spec_xlsx(est, number)
    files = [contract_out] + ([app1_out] if app1_out else [])
    return {"contract_no": number, "files": [f.name for f in files if f], "download": [f"/files?path={f.name}" for f in files if f]}

@router.post("/create-all")
def create_all(customer_id: int = Query(...), estimate_id: int | None = Query(None), kp_id: int | None = Query(None), db: Session = Depends(get_db)):
    number, ctx, est, kp = build_ctx(db, customer_id, estimate_id, kp_id)
    out_files = []

    def ensure(name: str) -> Path:
        p = TEMPLATES_DIR / name
        if not p.is_file():
            raise HTTPException(400, f"Нет шаблона: {name} в storage/templates/contracts")
        return p

    # Договор
    contract_tpl = ensure("Договор.docx")
    contract_out = DOCS_DIR / f"Договор_{number}.docx"
    render_docx(str(contract_tpl), str(contract_out), ctx)
    out_files.append(contract_out)

    # Прил. 1 (спецификация) — генерим из сметы/КП
    if est:
        from app.services.spec_builder import build_spec_docx
        from app.services.export_xlsx import export_spec_xlsx
        spec_out = Path(build_spec_docx(est, number)); out_files.append(spec_out)
        xlsx = Path(export_spec_xlsx(est, number)); out_files.append(xlsx)

    # Прил. 2–6 — по вашим шаблонам (подстановка [[...]] если есть)
    mapping = ctx
    mapping.setdefault("DOG_NO", ctx["CONTRACT_NO"])
    mapping.setdefault("DOG_DATE", ctx["CONTRACT_DATE"])
    for src_name, out_base in [
        ("Приложение_2_Проектная документация.docx", f"Приложение_2_Проектная документация_{number}.docx"),
        ("Приложение_3_График выполнения работ.docx", f"Приложение_3_График_{number}.docx"),
        ("Приложение_4_Акт приема-передачи работ.docx", f"Приложение_4_Акт приема-передачи_{number}.docx"),
        ("Приложение_5_Акт пусконаладочных работ.docx", f"Приложение_5_Акт ПНР_{number}.docx"),
        ("Приложение_6_Акт опрессовки и вакуумирования.docx", f"Приложение_6_Акт опрессовки_{number}.docx"),
    ]:
        src = TEMPLATES_DIR / src_name
        if src.is_file():
            dst = DOCS_DIR / out_base
            render_or_copy(src, dst, mapping)
            out_files.append(dst)

    return {
        "contract_no": number,
        "files": [p.name for p in out_files],
        "download": [f"/files?path={p.name}" for p in out_files],
    }
""")

# 7.3) Веб-обёртка для create-all (кнопка в UI)
def add_ui_handler_create_all(s):
    if '@app.get("/contracts/create-all"' in s:
        return s
    s += textwrap.dedent("""
    @app.get("/contracts/create-all")
    def ui_create_all(customer_id: int, estimate_id: int | None = None, kp_id: int | None = None):
        # просто прокси в API и редирект в список документов
        from app.api.routers.contracts import create_all
        create_all(customer_id=customer_id, estimate_id=estimate_id, kp_id=kp_id)  # создадим пакет
        return RedirectResponse(url="/documents", status_code=303)
    """)
    return s

patch_file("app/main.py", add_ui_handler_create_all)

# 8) Страница /proposals/new уже заменена; нужно чтобы POST создавал клиента при отсутствии id — у нас это уже есть

# 9) Подправим /templates — добавим подсказки по именованию файлов
w("app/web/templates/templates.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Шаблоны документов (DOCX)</h2>
<p>Загрузите шаблоны в <code>storage/templates/contracts</code> или используйте форму ниже.<br>
Имена файлов (рекомендуемые):</p>
<ul>
  <li>Договор.docx</li>
  <li>Приложение_2_Проектная документация.docx</li>
  <li>Приложение_3_График выполнения работ.docx</li>
  <li>Приложение_4_Акт приема-передачи работ.docx</li>
  <li>Приложение_5_Акт пусконаладочных работ.docx</li>
  <li>Приложение_6_Акт опрессовки и вакуумирования.docx</li>
</ul>
<p>Поддерживаются плейсхолдеры в виде <code>[[CONTRACT_NO]]</code>, <code>[[CONTRACT_DATE]]</code>, <code>[[CUSTOMER_FIO]]</code>, <code>[[OBJECT_ADDR]]</code> и др.</p>
<form class="card" method="post" action="/templates/upload" enctype="multipart/form-data">
  <input name="name" placeholder="Имя файла, например: Договор.docx" required>
  <input type="file" name="file" required>
  <button>Загрузить</button>
</form>
{% endblock %}
""")

print("== DONE ==")
print("Перезапустите run_kp.bat и проверьте:")
print(" - /proposals/new — создание КП (выбор клиента из списка или новый)")
print(" - Карточка клиента — расширенные поля и кнопка «Сформировать пакет (Договор + Прил.1–6)»")
print(" - /documents — появятся сформированные файлы")
