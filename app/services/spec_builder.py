from docx import Document
from app.config import DOCS_DIR
def fmt_money(x):
    try: return f"{float(x):,.0f}".replace(",", " ")
    except: return str(x)
def build_spec_docx(estimate, number: str) -> str:
    doc = Document()
    doc.add_heading(f"Приложение №1 к договору № {number}", level=1)
    doc.add_paragraph(f"Объект: {estimate.site_address or ''}")
    tbl = doc.add_table(rows=1, cols=8); hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text, hdr[5].text, hdr[6].text, hdr[7].text = (
        "Раздел","Поз.","Наименование","Ед.","Кол-во","Цена (мат.)","Цена (монтаж)","Сумма")
    sec_map = {s.id: s for s in estimate.sections}
    for it in estimate.items:
        r = tbl.add_row().cells; sec = sec_map.get(it.section_id)
        r[0].text = sec.title if sec else ""; r[1].text = it.position or ""; r[2].text = it.name
        r[3].text = it.unit or ""; r[4].text = str(it.qty or 0)
        r[5].text = fmt_money(it.price_material or 0); r[6].text = fmt_money(it.price_labor or 0)
        tot = (it.qty or 0)*((it.price_material or 0)+(it.price_labor or 0)); r[7].text = fmt_money(tot)
    out = DOCS_DIR / f"Приложение_1_Спецификация_{number}.docx"; doc.save(out); return str(out)
