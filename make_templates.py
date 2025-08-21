# Создаёт шаблон договора с плейсхолдерами и пример КП (DOCX) в проекте.
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

BASE = Path(__file__).resolve().parent
TPL_DIR = BASE / "storage" / "templates" / "contracts"
DOCS_DIR = BASE / "storage" / "docs"
TPL_DIR.mkdir(parents=True, exist_ok=True)
DOCS_DIR.mkdir(parents=True, exist_ok=True)

def heading(doc, text, size=16, bold=True, align="center"):
    p = doc.add_paragraph()
    run = p.add_run(text); run.bold = bold; run.font.size = Pt(size)
    p.alignment = {"center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                   "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                   "right": WD_PARAGRAPH_ALIGNMENT.RIGHT}[align]

def kv(doc, label, value):
    p = doc.add_paragraph(); r1 = p.add_run(label + " "); r1.bold = True; p.add_run(value)

# ---------- Договор ----------
dog = Document()
heading(dog, "ДОГОВОР № [[CONTRACT_NO]]", 18, True, "center")
heading(dog, "от [[CONTRACT_DATE]] г.", 12, False, "center")
dog.add_paragraph()
kv(dog, "Заказчик:", "[[CUSTOMER_FIO]]")
kv(dog, "Паспорт:", "[[CUSTOMER_PASSPORT]]")
kv(dog, "Адрес регистрации:", "[[CUSTOMER_ADDR_REG]]")
kv(dog, "Телефон:", "[[CUSTOMER_PHONE]]")
kv(dog, "Email:", "[[CUSTOMER_EMAIL]]")
dog.add_paragraph()
kv(dog, "Исполнитель:", "[[EXECUTOR_FIO]]")
kv(dog, "Паспорт/реквизиты:", "[[EXECUTOR_PASSPORT]]")
kv(dog, "Адрес:", "[[EXECUTOR_ADDR]]")
kv(dog, "Телефон:", "[[EXECUTOR_PHONE]]")
kv(dog, "Email:", "[[EXECUTOR_EMAIL]]")
kv(dog, "ИНН:", "[[EXECUTOR_INN]]")
kv(dog, "ОГРН:", "[[EXECUTOR_OGRN]]")
dog.add_paragraph(); kv(dog, "Объект:", "[[OBJECT_ADDR]]")

dog.add_paragraph()
heading(dog, "1. Стоимость работ и материалов", 14, True, "left")
t = dog.add_table(rows=4, cols=2)
t.rows[0].cells[0].text = "Позиция";             t.rows[0].cells[1].text = "Сумма, руб."
t.rows[1].cells[0].text = "Оборудование и материалы"; t.rows[1].cells[1].text = "[[SUM_EQUIP]]"
t.rows[2].cells[0].text = "Монтажные работы";    t.rows[2].cells[1].text = "[[SUM_WORK]]"
t.rows[3].cells[0].text = "ИТОГО";               t.rows[3].cells[1].text = "[[SUM_TOTAL]]"

dog.add_paragraph()
heading(dog, "2. Порядок расчётов", 14, True, "left")
dog.add_paragraph("2.1. Предоплата составляет [[PREPAYMENT_PCT]]% от общей стоимости.")
pt = dog.add_table(rows=4, cols=2)
pt.rows[0].cells[0].text = "Этап";   pt.rows[0].cells[1].text = "Сумма, руб."
pt.rows[1].cells[0].text = "Этап 1"; pt.rows[1].cells[1].text = "[[STAGE1]]"
pt.rows[2].cells[0].text = "Этап 2"; pt.rows[2].cells[1].text = "[[STAGE2]]"
pt.rows[3].cells[0].text = "Этап 3"; pt.rows[3].cells[1].text = "[[STAGE3]]"

dog.add_paragraph()
heading(dog, "3. Сроки и гарантия", 14, True, "left")
dog.add_paragraph("3.1. Срок поставки: [[DELIVERY_DAYS]] календарных дней.")
dog.add_paragraph("3.2. Срок монтажа: [[INSTALL_DAYS]] календарных дней.")
dog.add_paragraph("3.3. Гарантия: [[WARRANTY_MONTHS]] месяцев.")
dog.add_paragraph("3.4. Срок действия КП: [[KP_VALID_DAYS]] дней.")

dog.add_paragraph()
heading(dog, "4. Прочие условия", 14, True, "left")
dog.add_paragraph("4.1. Стороны обязуются выполнять условия настоящего договора добросовестно.")
dog.add_paragraph("4.2. Споры разрешаются путём переговоров.")

dog.add_paragraph()
heading(dog, "5. Реквизиты и подписи сторон", 14, True, "left")
st = dog.add_table(rows=2, cols=2)
st.cell(0,0).text = "Заказчик: [[CUSTOMER_FIO]]\\n_________________/________________"
st.cell(0,1).text = "Исполнитель: [[EXECUTOR_FIO]]\\n_________________/________________"
st.cell(1,0).text = "Тел.: [[CUSTOMER_PHONE]]  Email: [[CUSTOMER_EMAIL]]"
st.cell(1,1).text = "Тел.: [[EXECUTOR_PHONE]]  Email: [[EXECUTOR_EMAIL]]"

dog_out = TPL_DIR / "Договор.docx"
dog.save(dog_out)

# ---------- Пример КП (вне системы, просто образец) ----------
kp = Document()
def add_kv(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(label + " "); r1.bold = True
    p.add_run(value)

heading(kp, "VENT•PRO — HVAC Solutions", 16, True, "left")
heading(kp, "Коммерческое предложение", 18, True, "center")
add_kv(kp, "Клиент:", "ООО «Пример», ИНН 7700000000")
add_kv(kp, "Контакт:", "Иванов Иван Иванович, +7 (900) 000-00-00, example@mail.ru")
add_kv(kp, "Объект:", "г. Москва, ул. Примерная, д. 1, корп. 2")

kp.add_paragraph()
heading(kp, "Состав работ и оборудования", 14, True, "left")

tbl = kp.add_table(rows=1, cols=7)
hdr = tbl.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "Раздел", "Поз.", "Наименование"
hdr[3].text, hdr[4].text, hdr[5].text, hdr[6].text = "Ед.", "Кол-во", "Цена", "Сумма"

def add(razdel, poz, name, unit, qty, price):
    r = tbl.add_row().cells
    r[0].text, r[1].text, r[2].text = razdel, poz, name
    r[3].text, r[4].text = unit, str(qty)
    r[5].text = f"{price:,.0f}".replace(",", " ")
    r[6].text = f"{qty*price:,.0f}".replace(",", " ")

add("1. Вентиляция", "1.1", "Вентилятор осевой 250 мм", "шт", 2, 12500)
add("1. Вентиляция", "1.2", "Канал воздуховод оцинкованный Ø250", "м", 15, 800)
add("2. Кондиционирование", "2.1", "Сплит-система 3,5 кВт (LG)", "шт", 1, 52000)
add("2. Кондиционирование", "2.2", "Монтаж сплит-системы", "компл", 1, 15000)

items = [(2,12500),(15,800),(1,52000),(1,15000)]
total = sum(q*p for q,p in items)
p = kp.add_paragraph(); p.add_run("ИТОГО по КП: ").bold = True; p.add_run(f"{total:,.0f} руб.").bold = True
kp.add_paragraph("Срок действия предложения: 14 календарных дней.")
kp.add_paragraph("Срок поставки: 10 дней. Срок монтажа: 5 дней.")
kp.add_paragraph("Гарантия на оборудование: 24 месяца.")

kp_out = DOCS_DIR / "КП_шаблон_с_примером.docx"
kp.save(kp_out)

print("OK:", dog_out, "и", kp_out)
