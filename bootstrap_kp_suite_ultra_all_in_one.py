# -*- coding: utf-8 -*-
# Создаёт ПОЛНЫЙ проект "kp_suite_ultra" со всем кодом, шаблонами и скриптами.
# Запуск (Windows):  py -3 bootstrap_kp_suite_ultra_all_in_one.py

import os, textwrap
from pathlib import Path

ROOT = Path.cwd() / "kp_suite_ultra"

def w(relpath: str, content: str = ""):
    p = ROOT / relpath
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(textwrap.dedent(content).lstrip("\n"), encoding="utf-8")
    print("wrote:", relpath)

print("Creating project at:", ROOT)
(ROOT).mkdir(parents=True, exist_ok=True)

# ---------- top-level ----------
w("requirements.txt", """
fastapi
uvicorn
sqlalchemy>=2.0
pydantic>=2
python-multipart
Jinja2
starlette
openpyxl
python-docx
xhtml2pdf
""")

w("run_kp.bat", r"""
@echo off
setlocal enabledelayedexpansion

REM ===== KP Suite ULTRA launcher (robust) =====
cd /d %~dp0

if not exist .venv (
  echo Creating virtual environment...
  py -3 -m venv .venv
)

set "PY=.venv\Scripts\python.exe"

echo Activating venv...
call .venv\Scripts\activate

echo Upgrading pip inside venv...
"%PY%" -m pip install --upgrade pip

echo Installing requirements into venv...
"%PY%" -m pip install -r requirements.txt

echo Starting KP Suite ULTRA...
set ADMIN_USER=admin
set ADMIN_PASS=admin
"%PY%" -m uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload
""")

# ---------- app skeleton ----------
for d in [
  "app/db", "app/security", "app/models", "app/schemas", "app/services",
  "app/api/routers", "app/web/static", "app/web/templates",
  "storage/templates/contracts", "storage/docs"
]:
    (ROOT/d).mkdir(parents=True, exist_ok=True)

w("app/__init__.py", "")
w("app/config.py", """
import os
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent.parent
DB_URL = os.getenv("DB_URL", f"sqlite:///{(BASE_DIR / 'kp.db').as_posix()}")

DOCS_DIR = BASE_DIR / "storage" / "docs"
TEMPLATES_DIR = BASE_DIR / "storage" / "templates" / "contracts"
STATIC_DIR = BASE_DIR / "app" / "web" / "static"
HTML_TPL_DIR = BASE_DIR / "app" / "web" / "templates"

for p in (DOCS_DIR, TEMPLATES_DIR, STATIC_DIR, HTML_TPL_DIR):
    p.mkdir(parents=True, exist_ok=True)

CONTRACT_NUMBER_MASK = os.getenv("CONTRACT_NUMBER_MASK", "DOG-{YYYY}-{SEQ:04d}")
PROPOSAL_NUMBER_MASK = os.getenv("PROPOSAL_NUMBER_MASK", "KP-{YYYY}-{SEQ:04d}")

ADMIN_USER = os.getenv("ADMIN_USER", "admin")
ADMIN_PASS = os.getenv("ADMIN_PASS", "admin")
""")

# ---------- DB ----------
w("app/db/base.py", "from sqlalchemy.orm import declarative_base\nBase = declarative_base()\n")
w("app/db/session.py", """
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from app.config import DB_URL

connect_args = {"check_same_thread": False} if DB_URL.startswith("sqlite") else {}
engine = create_engine(DB_URL, connect_args=connect_args)
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
""")
w("app/db/init_db.py", """
from sqlalchemy import inspect
from app.db.session import engine
from app.db.base import Base
from app import models  # noqa

def init_db():
    Base.metadata.create_all(bind=engine)
    print("Tables:", inspect(engine).get_table_names())
""")

# ---------- security ----------
w("app/security/__init__.py", "")
w("app/security/basic.py", """
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from fastapi import Depends, HTTPException, status
from app.config import ADMIN_USER, ADMIN_PASS
import secrets

security = HTTPBasic()

def admin_guard(credentials: HTTPBasicCredentials = Depends(security)):
    correct_username = secrets.compare_digest(credentials.username, ADMIN_USER)
    correct_password = secrets.compare_digest(credentials.password, ADMIN_PASS)
    if not (correct_username and correct_password):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Unauthorized", headers={"WWW-Authenticate": "Basic"})
    return True
""")

# ---------- models ----------
w("app/models/__init__.py", """
from .sequence import Sequence
from .owner import Owner
from .customer import Customer
from .category import Category
from .material import Material
from .estimate import Estimate, EstimateSection, EstimateItem
from .proposal import Proposal, ProposalSection, ProposalItem
from .document import DocumentRec
from .template import TemplateRec
from .contract import ContractData
""")
w("app/models/sequence.py", """
from sqlalchemy import Column, Integer, String
from app.db.base import Base
class Sequence(Base):
    __tablename__ = "sequences"
    id = Column(Integer, primary_key=True)
    name = Column(String(50), unique=True, nullable=False)
    last = Column(Integer, default=0, nullable=False)
""")
w("app/models/owner.py", """
from sqlalchemy import Column, Integer, String
from app.db.base import Base
class Owner(Base):
    __tablename__ = "owner"
    id = Column(Integer, primary_key=True)
    display_name = Column(String(200), default="Исполнитель")
    fullname = Column(String(200))
    passport = Column(String(200))
    address = Column(String(300))
    phone = Column(String(100))
    email = Column(String(150))
    inn = Column(String(50))
    ogrn = Column(String(50))
""")
w("app/models/customer.py", """
from sqlalchemy import Column, Integer, String
from app.db.base import Base
class Customer(Base):
    __tablename__ = "customers"
    id = Column(Integer, primary_key=True)
    fullname = Column(String(200), nullable=False)
    passport = Column(String(200))
    address = Column(String(300))
    address_object = Column(String(300))
    phone = Column(String(100))
    email = Column(String(150))
""")
w("app/models/category.py", """
from sqlalchemy import Column, Integer, String
from app.db.base import Base
class Category(Base):
    __tablename__ = "categories"
    id = Column(Integer, primary_key=True)
    name = Column(String(200), nullable=False)
    parent_id = Column(Integer)
    kind = Column(String(50), nullable=False, default="materials")
""")
w("app/models/material.py", """
from sqlalchemy import Column, Integer, String, Float, ForeignKey
from app.db.base import Base
class Material(Base):
    __tablename__ = "materials"
    id = Column(Integer, primary_key=True)
    name = Column(String(300), nullable=False)
    model = Column(String(200))
    unit = Column(String(50), default="шт")
    price_material = Column(Float, default=0.0)
    price_labor = Column(Float, default=0.0)
    category_id = Column(Integer, ForeignKey("categories.id"), nullable=True)
""")
w("app/models/estimate.py", """
from sqlalchemy import Column, Integer, String, Float, ForeignKey
from sqlalchemy.orm import relationship
from app.db.base import Base

class Estimate(Base):
    __tablename__ = "estimates"
    id = Column(Integer, primary_key=True)
    title = Column(String(200), default="Смета")
    customer_id = Column(Integer, ForeignKey("customers.id"))
    site_address = Column(String(300))
    customer = relationship("Customer")

class EstimateSection(Base):
    __tablename__ = "estimate_sections"
    id = Column(Integer, primary_key=True)
    estimate_id = Column(Integer, ForeignKey("estimates.id"), nullable=False)
    title = Column(String(200), nullable=False)
    position = Column(String(20))
    estimate = relationship("Estimate", backref="sections")

class EstimateItem(Base):
    __tablename__ = "estimate_items"
    id = Column(Integer, primary_key=True)
    estimate_id = Column(Integer, ForeignKey("estimates.id"), nullable=False)
    section_id = Column(Integer, ForeignKey("estimate_sections.id"))
    name = Column(String(300), nullable=False)
    item_type = Column(String(50), default="material")
    unit = Column(String(50), default="шт")
    qty = Column(Float, default=1.0)
    price_material = Column(Float, default=0.0)
    price_labor = Column(Float, default=0.0)
    position = Column(String(20))
    estimate = relationship("Estimate", backref="items")
    section = relationship("EstimateSection")
""")
w("app/models/proposal.py", """
from sqlalchemy import Column, Integer, String, Float, ForeignKey, DateTime, func
from sqlalchemy.orm import relationship
from app.db.base import Base

class Proposal(Base):
    __tablename__ = "proposals"
    id = Column(Integer, primary_key=True)
    number = Column(String(50))
    title = Column(String(200), default="Коммерческое предложение")
    customer_id = Column(Integer, ForeignKey("customers.id"))
    site_address = Column(String(300))
    header_html = Column(String(500), default="")
    status = Column(String(50), default="draft")
    created_at = Column(DateTime, server_default=func.now())
    customer = relationship("Customer")

class ProposalSection(Base):
    __tablename__ = "proposal_sections"
    id = Column(Integer, primary_key=True)
    proposal_id = Column(Integer, ForeignKey("proposals.id"), nullable=False)
    title = Column(String(200), nullable=False)
    position = Column(String(20))
    proposal = relationship("Proposal", backref="sections")

class ProposalItem(Base):
    __tablename__ = "proposal_items"
    id = Column(Integer, primary_key=True)
    proposal_id = Column(Integer, ForeignKey("proposals.id"), nullable=False)
    section_id = Column(Integer, ForeignKey("proposal_sections.id"))
    name = Column(String(300), nullable=False)    # Наименование и тех. характеристика
    note = Column(String(300))                    # Тип, марка, обозначение
    unit = Column(String(50), default="шт")
    qty = Column(Float, default=1.0)
    price = Column(Float, default=0.0)            # материалы/оборудование, за ед.
    price_labor = Column(Float, default=0.0)      # монтаж, за ед.
    position = Column(String(20))
    proposal = relationship("Proposal", backref="items")
    section = relationship("ProposalSection")
""")
w("app/models/document.py", """
from sqlalchemy import Column, Integer, String, DateTime, func
from app.db.base import Base
class DocumentRec(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True)
    customer_id = Column(Integer)
    estimate_id = Column(Integer)
    kind = Column(String(50))  # contract, kp, appendix1..7, spec, etc.
    title = Column(String(200))
    number = Column(String(100))
    filepath = Column(String(500), nullable=False)
    status = Column(String(50), default="draft")
    created_at = Column(DateTime, server_default=func.now())
""")
w("app/models/template.py", """
from sqlalchemy import Column, Integer, String, DateTime, func
from app.db.base import Base
class TemplateRec(Base):
    __tablename__ = "templates"
    id = Column(Integer, primary_key=True)
    name = Column(String(200), unique=True)  # 'Договор.docx', ...
    path = Column(String(500), nullable=False)
    uploaded_at = Column(DateTime, server_default=func.now())
""")
w("app/models/contract.py", """
from sqlalchemy import Column, Integer, String, DateTime, func, Text
from app.db.base import Base
class ContractData(Base):
    __tablename__ = "contract_data"
    id = Column(Integer, primary_key=True)
    number = Column(String(50))
    customer_id = Column(Integer)
    estimate_id = Column(Integer)
    kp_id = Column(Integer)
    ctx_json = Column(Text)  # сохраненные поля договора
    created_at = Column(DateTime, server_default=func.now())
""")

# ---------- schemas ----------
w("app/schemas/__init__.py", "")
w("app/schemas/customer.py", """
from pydantic import BaseModel, Field
class CustomerIn(BaseModel):
    fullname: str = Field(...)
    passport: str | None = None
    address: str | None = None
    address_object: str | None = None
    phone: str | None = None
    email: str | None = None
class CustomerOut(CustomerIn):
    id: int
    class Config: from_attributes = True
""")
w("app/schemas/estimate.py", """
from pydantic import BaseModel
from typing import Optional
class EstimateIn(BaseModel):
    title: str = "Смета"
    customer_id: int
    site_address: Optional[str] = None
class EstimateOut(EstimateIn):
    id: int
    class Config: from_attributes = True
class SectionIn(BaseModel):
    estimate_id: int
    title: str
class SectionOut(SectionIn):
    id: int
    position: str | None = None
    class Config: from_attributes = True
class ItemIn(BaseModel):
    estimate_id: int
    section_id: int | None = None
    name: str
    item_type: str = "material"
    unit: str = "шт"
    qty: float = 1.0
    price_material: float = 0.0
    price_labor: float = 0.0
class ItemOut(ItemIn):
    id: int
    position: str | None = None
    class Config: from_attributes = True
""")
w("app/schemas/proposal.py", """
from pydantic import BaseModel
from typing import Optional
class ProposalIn(BaseModel):
    title: str = "Коммерческое предложение"
    customer_id: int
    site_address: Optional[str] = None
    header_html: Optional[str] = ""
class ProposalOut(ProposalIn):
    id: int
    number: str | None = None
    class Config: from_attributes = True
class PSectionIn(BaseModel):
    proposal_id: int
    title: str
class PSectionOut(PSectionIn):
    id: int
    position: str | None = None
    class Config: from_attributes = True
class PItemIn(BaseModel):
    proposal_id: int
    section_id: int | None = None
    name: str
    note: str | None = None
    unit: str = "шт"
    qty: float = 1.0
    price: float = 0.0
    price_labor: float = 0.0
class PItemOut(PItemIn):
    id: int
    position: str | None = None
    class Config: from_attributes = True
""")
w("app/schemas/category.py", """
from pydantic import BaseModel
class CategoryIn(BaseModel):
    name: str
    parent_id: int | None = None
    kind: str = "materials"
class CategoryOut(CategoryIn):
    id: int
    class Config: from_attributes = True
""")
w("app/schemas/material.py", """
from pydantic import BaseModel
class MaterialIn(BaseModel):
    name: str
    model: str | None = None
    unit: str = "шт"
    price_material: float = 0.0
    price_labor: float = 0.0
    category_id: int | None = None
class MaterialOut(MaterialIn):
    id: int
    class Config: from_attributes = True
""")

# ---------- services ----------
w("app/services/__init__.py", "")
w("app/services/numbering.py", """
from datetime import datetime
from sqlalchemy.orm import Session
from sqlalchemy import select
from app.models import Sequence
from app.config import CONTRACT_NUMBER_MASK, PROPOSAL_NUMBER_MASK

def next_seq(db: Session, name: str) -> int:
    row = db.execute(select(Sequence).where(Sequence.name==name)).scalar_one_or_none()
    if not row:
        row = Sequence(name=name, last=0); db.add(row); db.commit(); db.refresh(row)
    row.last += 1; db.add(row); db.commit(); db.refresh(row); return row.last

def generate_contract_number(db: Session) -> str:
    now = datetime.now(); seq = next_seq(db, "DOG")
    return CONTRACT_NUMBER_MASK.replace("{YYYY}", f"{now.year}").replace("{SEQ:04d}", f"{seq:04d}")

def generate_proposal_number(db: Session) -> str:
    now = datetime.now(); seq = next_seq(db, "KP")
    return PROPOSAL_NUMBER_MASK.replace("{YYYY}", f"{now.year}").replace("{SEQ:04d}", f"{seq:04d}")
""")
w("app/services/payments.py", """
def estimate_totals(estimate) -> dict:
    equip = sum((i.qty or 0)*(i.price_material or 0) for i in estimate.items)
    labor = sum((i.qty or 0)*(i.price_labor or 0) for i in estimate.items)
    return {"equipment": equip, "labor": labor, "total": equip + labor}

def proposal_totals(proposal) -> dict:
    equip = sum((i.qty or 0)*(i.price or 0) for i in proposal.items)
    labor = sum((i.qty or 0)*(i.price_labor or 0) for i in proposal.items)
    return {"equipment": equip, "labor": labor, "total": equip + labor}

def payment_schedule_100_70_30_from_totals(equip: float, labor: float) -> list[dict]:
    return [
        {"title": "100% оборудования и материалов", "amount": equip},
        {"title": "70% монтажных работ", "amount": labor*0.7},
        {"title": "30% монтажных работ", "amount": labor*0.3},
    ]

def payment_schedule_100_70_30(estimate) -> list[dict]:
    t = estimate_totals(estimate)
    return payment_schedule_100_70_30_from_totals(t["equipment"], t["labor"])
""")
w("app/services/files.py", """
from pathlib import Path
from fastapi import HTTPException
from fastapi.responses import FileResponse
from app.config import DOCS_DIR
SAFE_ROOT = DOCS_DIR.resolve()
def safe_file_response(rel_or_abs_path: str):
    p = Path(rel_or_abs_path)
    p = (SAFE_ROOT / p).resolve() if not p.is_absolute() else p.resolve()
    if SAFE_ROOT not in p.parents and p != SAFE_ROOT:
        raise HTTPException(403, "forbidden")
    if not p.is_file():
        raise HTTPException(404, "file not found")
    return FileResponse(str(p))
""")
w("app/services/docx_merge.py", """
from docx import Document
from pathlib import Path
import os
def _replace_in_par(par, mapping: dict):
    text = "".join(r.text for r in par.runs); changed = False
    for k, v in mapping.items():
        ph = f"[[{k}]]"
        if ph in text: text = text.replace(ph, str(v)); changed = True
    if changed:
        for _ in range(len(par.runs)-1): par.runs[-1].clear()
        par.runs[0].text = text
def _walk(cell_or_doc, mapping: dict):
    for p in getattr(cell_or_doc, "paragraphs", []): _replace_in_par(p, mapping)
    for t in getattr(cell_or_doc, "tables", []):
        for row in t.rows:
            for cell in row.cells: _walk(cell, mapping)
def render_docx(src_path: str, dst_path: str, mapping: dict) -> str:
    doc = Document(src_path); _walk(doc, mapping)
    Path(os.path.dirname(dst_path)).mkdir(parents=True, exist_ok=True)
    doc.save(dst_path); return dst_path
""")
w("app/services/spec_builder.py", """
from docx import Document
from app.config import DOCS_DIR
def fmt_money(x):
    try: return f"{float(x):,.0f}".replace(",", " ")
    except: return str(x)
def build_spec_docx(estimate, number: str) -> str:
    doc = Document()
    doc.add_heading(f"Приложение №1 к договору № {number}", level=1)
    doc.add_paragraph(f"Объект: {estimate.site_address or ''}")
    tbl = doc.add_table(rows=1, cols=8); hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text, hdr[5].text, hdr[6].text, hdr[7].text = (
        "Раздел","Поз.","Наименование","Ед.","Кол-во","Цена (мат.)","Цена (монтаж)","Сумма")
    sec_map = {s.id: s for s in estimate.sections}
    for it in estimate.items:
        r = tbl.add_row().cells; sec = sec_map.get(it.section_id)
        r[0].text = sec.title if sec else ""; r[1].text = it.position or ""; r[2].text = it.name
        r[3].text = it.unit or ""; r[4].text = str(it.qty or 0)
        r[5].text = fmt_money(it.price_material or 0); r[6].text = fmt_money(it.price_labor or 0)
        tot = (it.qty or 0)*((it.price_material or 0)+(it.price_labor or 0)); r[7].text = fmt_money(tot)
    out = DOCS_DIR / f"Приложение_1_Спецификация_{number}.docx"; doc.save(out); return str(out)
""")
w("app/services/export_xlsx.py", """
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from app.config import DOCS_DIR
def export_spec_xlsx(estimate, number: str) -> str:
    wb = Workbook(); ws = wb.active; ws.title = "Спецификация"
    ws.append(["Раздел","Поз.","Наименование","Ед.","Кол-во","Цена (мат.)","Цена (монтаж)","Сумма"])
    sec_map = {s.id: s for s in estimate.sections}
    for it in estimate.items:
        sec = sec_map.get(it.section_id); total = (it.qty or 0)*((it.price_material or 0)+(it.price_labor or 0))
        ws.append([sec.title if sec else "", it.position or "", it.name, it.unit or "", it.qty or 0, it.price_material or 0, it.price_labor or 0, total])
    for c in range(1,9): ws.column_dimensions[get_column_letter(c)].width = 18
    path = DOCS_DIR / f"Спецификация_{number}.xlsx"; wb.save(path); return str(path)
""")
w("app/services/ai.py", """
from difflib import get_close_matches
DEFAULT_SECTIONS = [
    "Оборудование", "Вентиляция", "Кондиционирование",
    "Электрика", "Сантехника", "Изоляция", "Управление и автоматика"
]
KEYWORDS2SECTION = {
    "вентилятор": "Вентиляция", "канал": "Вентиляция", "решетка": "Вентиляция",
    "lg": "Кондиционирование", "vrf": "Кондиционирование",
    "кабель": "Электрика", "автомат": "Электрика",
    "насос": "Сантехника", "изоляц": "Изоляция",
    "контроллер": "Управление и автоматика", "датчик": "Управление и автоматика",
}
def suggest_sections_from_text(text: str) -> list[str]:
    text_l = text.lower(); found = {sec for k, sec in KEYWORDS2SECTION.items() if k in text_l}
    return list(found) or DEFAULT_SECTIONS[:3]
def fuzzy_category(name: str, catalog: list[str]) -> str | None:
    m = get_close_matches(name, catalog, n=1, cutoff=0.6); return m[0] if m else None
""")
w("app/services/csv_import.py", """
import csv, io
from app.models import Material

def import_materials_csv(db, content: bytes):
    # Try utf-8 then cp1251
    try:
        text = content.decode('utf-8')
    except UnicodeDecodeError:
        text = content.decode('cp1251', errors='ignore')
    reader = csv.DictReader(io.StringIO(text))
    created = 0
    for row in reader:
        name = (row.get('name') or row.get('Наименование') or '').strip()
        if not name: continue
        m = Material(
            name=name,
            model=(row.get('model') or row.get('Модель') or '').strip() or None,
            unit=(row.get('unit') or row.get('Ед') or 'шт').strip() or 'шт',
            price_material=float((row.get('price_material') or row.get('ЦенаМат') or 0) or 0),
            price_labor=float((row.get('price_labor') or row.get('ЦенаМонтаж') or 0) or 0),
        )
        db.add(m); created += 1
    db.commit()
    return {"created": created}
""")

# ---------- API routers ----------
w("app/api/__init__.py", "")
w("app/api/deps.py", """
from fastapi import Depends
from app.db.session import SessionLocal
def get_db():
    db = SessionLocal()
    try: yield db
    finally: db.close()
""")

w("app/api/routers/__init__.py", "")
w("app/api/routers/customers.py", """
from fastapi import APIRouter, Depends
from sqlalchemy.orm import Session
from sqlalchemy import select
from app.api.deps import get_db
from app.models import Customer
from app.schemas.customer import CustomerIn, CustomerOut

router = APIRouter(prefix="/api/v1/customers", tags=["customers"])

@router.get("", response_model=list[CustomerOut])
def list_customers(db: Session = Depends(get_db)):
    return db.execute(select(Customer)).scalars().all()

@router.post("", response_model=CustomerOut)
def create_customer(payload: CustomerIn, db: Session = Depends(get_db)):
    c = Customer(**payload.dict()); db.add(c); db.commit(); db.refresh(c); return c
""")
w("app/api/routers/categories.py", """
from fastapi import APIRouter, Depends
from sqlalchemy.orm import Session
from sqlalchemy import select
from app.api.deps import get_db
from app.models import Category
from app.schemas.category import CategoryIn, CategoryOut

router = APIRouter(prefix="/api/v1/categories", tags=["categories"])

@router.get("", response_model=list[CategoryOut])
def list_categories(db: Session = Depends(get_db)):
    return db.execute(select(Category)).scalars().all()

@router.post("", response_model=CategoryOut)
def create_category(payload: CategoryIn, db: Session = Depends(get_db)):
    c = Category(**payload.dict()); db.add(c); db.commit(); db.refresh(c); return c
""")
w("app/api/routers/materials.py", """
from fastapi import APIRouter, Depends, UploadFile, File
from sqlalchemy.orm import Session
from sqlalchemy import select
from app.api.deps import get_db
from app.models import Material
from app.schemas.material import MaterialIn, MaterialOut
from app.services.csv_import import import_materials_csv

router = APIRouter(prefix="/api/v1/materials", tags=["materials"])

@router.get("", response_model=list[MaterialOut])
def list_materials(q: str | None = None, limit: int = 100, db: Session = Depends(get_db)):
    qset = select(Material)
    if q:
        like = f"%{q}%"
        qset = qset.where(Material.name.like(like) | Material.model.like(like))
    return db.execute(qset.limit(limit)).scalars().all()

@router.get("/suggest", response_model=list[MaterialOut])
def suggest_materials(q: str, db: Session = Depends(get_db)):
    like = f"%{q}%"
    qset = select(Material).where(Material.name.like(like) | Material.model.like(like)).limit(10)
    return db.execute(qset).scalars().all()

@router.post("", response_model=MaterialOut)
def create_material(payload: MaterialIn, db: Session = Depends(get_db)):
    m = Material(**payload.dict()); db.add(m); db.commit(); db.refresh(m); return m

@router.post("/import-csv")
async def import_csv(file: UploadFile = File(...), db: Session = Depends(get_db)):
    content = await file.read()
    result = import_materials_csv(db, content)
    return {"status": "ok", **result}
""")
w("app/api/routers/estimates.py", """
from fastapi import APIRouter, Depends
from sqlalchemy.orm import Session
from sqlalchemy import select, func
from app.api.deps import get_db
from app.models import Estimate, EstimateSection, EstimateItem
from app.schemas.estimate import EstimateIn, EstimateOut, SectionIn, SectionOut, ItemIn, ItemOut

router = APIRouter(prefix="/api/v1/estimates", tags=["estimates"])

@router.get("", response_model=list[EstimateOut])
def list_estimates(db: Session = Depends(get_db)):
    return db.execute(select(Estimate)).scalars().all()

@router.post("", response_model=EstimateOut)
def create_estimate(payload: EstimateIn, db: Session = Depends(get_db)):
    e = Estimate(**payload.dict()); db.add(e); db.commit(); db.refresh(e); return e

def _next_position_for_section(db: Session, section_id: int) -> str:
    sec = db.get(EstimateSection, section_id)
    if not sec: return ""
    try:
        sec_num = int((sec.position or "").split(".")[0])
    except:
        sections = db.execute(select(EstimateSection).where(EstimateSection.estimate_id==sec.estimate_id)).scalars().all()
        sec_num = sections.index(sec)+1
    count = db.execute(select(func.count(EstimateItem.id)).where(EstimateItem.section_id==section_id)).scalar_one()
    return f"{sec_num}.{count+1}"

@router.post("/sections", response_model=SectionOut)
def add_section(payload: SectionIn, db: Session = Depends(get_db)):
    count = db.execute(select(func.count(EstimateSection.id)).where(EstimateSection.estimate_id==payload.estimate_id)).scalar_one()
    s = EstimateSection(estimate_id=payload.estimate_id, title=payload.title, position=str(count+1))
    db.add(s); db.commit(); db.refresh(s); return s

@router.post("/items", response_model=ItemOut)
def add_item(payload: ItemIn, db: Session = Depends(get_db)):
    pos = _next_position_for_section(db, payload.section_id) if payload.section_id else None
    i = EstimateItem(**payload.dict(), position=pos); db.add(i); db.commit(); db.refresh(i); return i
""")
w("app/api/routers/proposals.py", """
from fastapi import APIRouter, Depends
from sqlalchemy.orm import Session
from sqlalchemy import select, func
from app.api.deps import get_db
from app.models import Proposal, ProposalSection, ProposalItem
from app.schemas.proposal import ProposalIn, ProposalOut, PSectionIn, PSectionOut, PItemIn, PItemOut
from app.services.numbering import generate_proposal_number

router = APIRouter(prefix="/api/v1/proposals", tags=["proposals"])

@router.get("", response_model=list[ProposalOut])
def list_proposals(db: Session = Depends(get_db)):
    return db.execute(select(Proposal).order_by(Proposal.created_at.desc())).scalars().all()

@router.post("", response_model=ProposalOut)
def create_proposal(payload: ProposalIn, db: Session = Depends(get_db)):
    number = generate_proposal_number(db)
    p = Proposal(**payload.dict(), number=number)
    db.add(p); db.commit(); db.refresh(p); return p

def _next_position_for_section(db: Session, section_id: int) -> str:
    sec = db.get(ProposalSection, section_id)
    if not sec: return ""
    try:
        sec_num = int((sec.position or "").split(".")[0])
    except:
        sections = db.execute(select(ProposalSection).where(ProposalSection.proposal_id==sec.proposal_id)).scalars().all()
        sec_num = sections.index(sec)+1
    count = db.execute(select(func.count(ProposalItem.id)).where(ProposalItem.section_id==section_id)).scalar_one()
    return f"{sec_num}.{count+1}"

@router.post("/sections", response_model=PSectionOut)
def add_section(payload: PSectionIn, db: Session = Depends(get_db)):
    count = db.execute(select(func.count(ProposalSection.id)).where(ProposalSection.proposal_id==payload.proposal_id)).scalar_one()
    s = ProposalSection(proposal_id=payload.proposal_id, title=payload.title, position=str(count+1))
    db.add(s); db.commit(); db.refresh(s); return s

@router.post("/items", response_model=PItemOut)
def add_item(payload: PItemIn, db: Session = Depends(get_db)):
    pos = _next_position_for_section(db, payload.section_id) if payload.section_id else None
    i = ProposalItem(**payload.dict(), position=pos); db.add(i); db.commit(); db.refresh(i); return i
""")
w("app/api/routers/contracts.py", """
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pathlib import Path
import json
from datetime import date
from app.api.deps import get_db
from app.models import Customer, Owner, Estimate, Proposal, DocumentRec, ContractData
from app.services.numbering import generate_contract_number
from app.services.payments import estimate_totals, payment_schedule_100_70_30, proposal_totals
from app.services.docx_merge import render_docx
from app.services.spec_builder import build_spec_docx
from app.services.export_xlsx import export_spec_xlsx
from app.config import TEMPLATES_DIR, DOCS_DIR

router = APIRouter(prefix="/api/v1/contracts", tags=["contracts"])

def _fmt(x):
    try: return f"{float(x):,.0f}".replace(",", " ")
    except: return str(x)

def build_ctx(db: Session, customer_id: int, estimate_id: int | None, kp_id: int | None, overrides: dict | None = None):
    cust = db.get(Customer, customer_id)
    if not cust: raise HTTPException(404, "customer not found")
    owner = db.get(Owner, 1)
    if not owner:
        owner = Owner(id=1, display_name="Исполнитель"); db.add(owner); db.commit(); db.refresh(owner)
    est = db.get(Estimate, estimate_id) if estimate_id else None
    kp = db.get(Proposal, kp_id) if kp_id else None
    totals = {"equipment":0,"labor":0,"total":0}; sched = []
    if est:
        totals = estimate_totals(est); sched = payment_schedule_100_70_30(est)
    elif kp:
        totals = proposal_totals(kp); sched = [{"title":"Оплата по КП","amount":totals["total"]}]
    number = generate_contract_number(db)
    ctx = {
        "CONTRACT_NO": number, "CONTRACT_DATE": date.today().strftime("%d.%m.%Y"),
        "CUSTOMER_FIO": cust.fullname or "", "CUSTOMER_PASSPORT": cust.passport or "",
        "CUSTOMER_ADDR_REG": cust.address or "", "CUSTOMER_PHONE": cust.phone or "",
        "CUSTOMER_EMAIL": cust.email or "",
        "OBJECT_ADDR": (getattr(est, "site_address", None) if est else None) or (kp.site_address if kp else None) or (cust.address_object or ""),
        "EXECUTOR_FIO": owner.fullname or owner.display_name or "Исполнитель",
        "EXECUTOR_PASSPORT": owner.passport or "", "EXECUTOR_ADDR": owner.address or "",
        "EXECUTOR_PHONE": owner.phone or "", "EXECUTOR_EMAIL": owner.email or "",
        "EXECUTOR_INN": owner.inn or "", "EXECUTOR_OGRN": owner.ogrn or "",
        "SUM_EQUIP": _fmt(totals["equipment"]), "SUM_WORK":  _fmt(totals["labor"]),
        "SUM_TOTAL": _fmt(totals["total"]),
        "STAGE1": _fmt(sched[0]["amount"] if sched else 0),
        "STAGE2": _fmt(sched[1]["amount"] if len(sched)>1 else 0),
        "STAGE3": _fmt(sched[2]["amount"] if len(sched)>2 else 0),
        "DELIVERY_DAYS": overrides.get("DELIVERY_DAYS") if overrides else "",
        "INSTALL_DAYS": overrides.get("INSTALL_DAYS") if overrides else "",
        "WARRANTY_MONTHS": overrides.get("WARRANTY_MONTHS") if overrides else "",
        "KP_VALID_DAYS": overrides.get("KP_VALID_DAYS") if overrides else "",
        "PREPAYMENT_PCT": overrides.get("PREPAYMENT_PCT") if overrides else "",
    }
    if overrides:
        ctx.update({k: v for k, v in overrides.items() if k in ctx})
    return number, ctx, est, kp

@router.post("/create")
def create_contract(customer_id: int = Query(...), estimate_id: int | None = Query(None), kp_id: int | None = Query(None), db: Session = Depends(get_db)):
    number, ctx, est, kp = build_ctx(db, customer_id, estimate_id, kp_id)
    dog_tpl = TEMPLATES_DIR / "Договор.docx"
    if not dog_tpl.is_file(): raise HTTPException(400, "Нет шаблона Договор.docx в storage/templates/contracts")
    contract_out = DOCS_DIR / f"Договор_{number}.docx"
    render_docx(str(dog_tpl), str(contract_out), ctx)
    app1_out = None
    if est:
        app1_out = Path(build_spec_docx(est, number)); export_spec_xlsx(est, number)
    files = [contract_out] + ([app1_out] if app1_out else [])
    return {"contract_no": number, "files": [f.name for f in files if f], "download": [f"/files?path={f.name}" for f in files if f]}

@router.post("/save-context")
def save_context(customer_id: int, estimate_id: int | None = None, kp_id: int | None = None, overrides: dict | None = None, db: Session = Depends(get_db)):
    number, ctx, _, _ = build_ctx(db, customer_id, estimate_id, kp_id, overrides or {})
    row = ContractData(number=number, customer_id=customer_id, estimate_id=estimate_id, kp_id=kp_id, ctx_json=json.dumps(ctx, ensure_ascii=False))
    db.add(row); db.commit(); return {"status":"ok", "number": number}
""")
w("app/api/routers/templates.py", """
from fastapi import APIRouter, Depends, UploadFile, File, Form
from sqlalchemy.orm import Session
from sqlalchemy import select
from pathlib import Path
from app.api.deps import get_db
from app.models import TemplateRec
from app.config import TEMPLATES_DIR
from app.security.basic import admin_guard

router = APIRouter(prefix="/api/v1/templates", tags=["templates"])

@router.get("")
def list_templates(db: Session = Depends(get_db)):
    rows = db.execute(select(TemplateRec)).scalars().all()
    return [{"id": r.id, "name": r.name, "path": Path(r.path).name, "uploaded_at": r.uploaded_at.isoformat()} for r in rows]

@router.post("/upload", dependencies=[Depends(admin_guard)])
async def upload_template(name: str = Form(...), file: UploadFile = File(...), db: Session = Depends(get_db)):
    path = TEMPLATES_DIR / name
    content = await file.read()
    path.write_bytes(content)
    row = db.execute(select(TemplateRec).where(TemplateRec.name==name)).scalar_one_or_none()
    if not row:
        row = TemplateRec(name=name, path=str(path)); db.add(row)
    else:
        row.path = str(path)
    db.commit()
    return {"status": "ok", "name": name}
""")
w("app/api/routers/files.py", """
from fastapi import APIRouter, Query
from app.services.files import safe_file_response
router = APIRouter(tags=["files"])
@router.get("/files")
def files(path: str = Query(..., description="Имя файла из storage/docs")):
    return safe_file_response(path)
""")
w("app/api/routers/aihelper.py", """
from fastapi import APIRouter
from app.services.ai import suggest_sections_from_text
router = APIRouter(prefix="/api/v1/ai", tags=["ai"])
@router.get("/suggest-sections")
def suggest_sections(text: str):
    return {"sections": suggest_sections_from_text(text)}
""")

# ---------- web (css + pages) ----------
w("app/web/static/main.css", """
:root{--bg:#0b0f14;--card:#0f172a;--border:#1f2937;--fg:#e8eef5;--accent:#a6e1ff}
body{font-family:system-ui,Arial,sans-serif;margin:0;background:var(--bg);color:var(--fg)}
header{background:#111826;padding:12px 16px;border-bottom:1px solid #1e293b;display:flex;justify-content:space-between;align-items:center}
header .logo{font-weight:800;letter-spacing:.5px}
header nav a{color:var(--accent);margin-right:12px;text-decoration:none}
main{padding:16px}
.card{background:var(--card);border:1px solid var(--border);border-radius:12px;padding:12px;margin-bottom:12px}
.grid{display:grid;grid-template-columns:1fr 1fr;gap:12px}
.grid2{display:grid;grid-template-columns:220px 1fr;gap:8px;align-items:center}
input,button,textarea{padding:8px;border-radius:8px;border:1px solid #334155;background:#0b1220;color:var(--fg)}
button{cursor:pointer}
.btn{padding:6px 10px;border:1px solid #334155;border-radius:8px;background:#0b1220;color:var(--fg);text-decoration:none}
.toolbar{display:flex;justify-content:space-between;align-items:center;background:#0f172a;border:1px solid var(--border);border-radius:12px;padding:8px;margin-bottom:12px;position:sticky;top:0;z-index:10}
.toolbar .primary{background:#0ea5e9;border-color:#0284c7}
table{width:100%;border-collapse:collapse}
th,td{border-bottom:1px solid var(--border);padding:8px}
.kp-head{font-weight:600}
""")
w("app/web/templates/layout.html", """
<!doctype html>
<html lang="ru">
<head>
  <meta charset="utf-8">
  <title>{{ title or "KP Suite ULTRA" }}</title>
  <link rel="stylesheet" href="/static/main.css">
</head>
<body>
<header>
  <div class="brand"><span class="logo">VENT•PRO</span> <span class="subtitle">HVAC / AI-помощник</span></div>
  <nav>
    <a href="/">Главная</a>
    <a href="/customers">Клиенты</a>
    <a href="/proposals">КП</a>
    <a href="/proposals/new">Быстрое КП</a>
    <a href="/materials">Материалы</a>
    <a href="/templates">Шаблоны</a>
    <a href="/documents">Документы</a>
  </nav>
</header>
<main>
  {% block content %}{% endblock %}
</main>
</body>
</html>
""")
w("app/web/templates/index.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Быстрый старт</h2>
<ol>
  <li>Клиенты → добавьте клиента (адрес регистрации и адрес объекта).</li>
  <li>КП → создайте КП из карточки клиента, добавляйте разделы/позиции (две цены: материал и монтаж) — автонумерация.</li>
  <li>Печать/В PDF — формат таблицы как в Excel. Внизу — этапы оплаты (100% / 70% / 30%).</li>
  <li>«Редактор договора» — подтянет суммы и адрес, создаст DOCX по шаблону.</li>
</ol>
{% endblock %}
""")
w("app/web/templates/customers.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Клиенты</h2>
<form class="card" method="post" action="/customers/create">
  <div class="grid2">
    <label>ФИО:</label><input name="fullname" placeholder="ФИО клиента" required>
    <label>Паспорт:</label><input name="passport" placeholder="Паспортные данные">
    <label>Телефон:</label><input name="phone" placeholder="+7 ...">
    <label>Email:</label><input name="email" placeholder="mail@example.com">
    <label>Адрес регистрации:</label><input name="address" placeholder="Город, улица ...">
    <label>Адрес объекта:</label><input name="address_object" placeholder="Адрес объекта (для КП/Договора)">
  </div>
  <button>Добавить клиента</button>
</form>

<table>
  <tr><th>ID</th><th>ФИО</th><th>Адрес объекта</th><th>Действия</th></tr>
  {% for c in customers %}
  <tr>
    <td>{{ c.id }}</td>
    <td><a href="/customers/card/{{ c.id }}">{{ c.fullname }}</a></td>
    <td>{{ c.address_object or "" }}</td>
    <td>
      <a class="btn" href="/proposals/create?customer_id={{ c.id }}">Создать КП</a>
      <a class="btn" href="/contract/editor?customer_id={{ c.id }}">Редактор договора</a>
    </td>
  </tr>
  {% endfor %}
</table>
{% endblock %}
""")
w("app/web/templates/customer_card_full.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Карточка клиента #{{ c.id }}</h2>
<div class="grid">
  <form class="card" method="post" action="/customers/{{ c.id }}/save">
    <h3>Основные данные</h3>
    <div class="grid2">
      <label>ФИО:</label><input name="fullname" value="{{ c.fullname or '' }}">
      <label>Паспорт:</label><input name="passport" value="{{ c.passport or '' }}">
      <label>Телефон:</label><input name="phone" value="{{ c.phone or '' }}">
      <label>Email:</label><input name="email" value="{{ c.email or '' }}">
      <label>Адрес регистрации:</label><input name="address" value="{{ c.address or '' }}">
      <label>Адрес объекта:</label><input name="address_object" value="{{ c.address_object or '' }}">
    </div>
    <div style="margin-top:8px">
      <button>Сохранить</button>
      <a class="btn" href="/proposals/create?customer_id={{ c.id }}">Создать КП</a>
      <a class="btn" href="/proposals/new?customer_id={{ c.id }}">Быстрое КП</a>
      <a class="btn" href="/contract/editor?customer_id={{ c.id }}{% if latest_kp %}&kp_id={{ latest_kp.id }}{% endif %}">Редактор договора{% if latest_kp %} (посл. КП {{ latest_kp.number }}){% endif %}</a>
    </div>
  </form>

  <div class="card">
    <h3>КП клиента</h3>
    <ul>
      {% for p in proposals %}
        <li><a href="/proposals/{{ p.id }}/edit">{{ p.created_at.strftime('%Y-%m-%d') if p.created_at else '' }} — {{ p.number }} — {{ p.title }}</a></li>
      {% else %}
        <li>КП пока нет</li>
      {% endfor %}
    </ul>
    <form method="post" action="/customers/{{ c.id }}/make-contract-latest-kp">
      <button>Сформировать договор (из последнего КП)</button>
    </form>
  </div>
</div>
{% endblock %}
""")
w("app/web/templates/proposals.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Коммерческие предложения</h2>
<p><a class="btn" href="/proposals/new">Быстрое создание КП</a></p>
<table>
  <tr><th>Дата</th><th>Номер</th><th>Клиент</th><th>Адрес</th><th>Открыть</th></tr>
  {% for p in proposals %}
  <tr>
    <td>{{ p.created_at.strftime('%Y-%m-%d') if p.created_at else '' }}</td>
    <td>{{ p.number }}</td>
    <td>{{ p.customer.fullname }}</td>
    <td>{{ p.site_address or p.customer.address_object or '' }}</td>
    <td><a href="/proposals/{{ p.id }}/edit">Редактировать</a></td>
  </tr>
  {% endfor %}
</table>
{% endblock %}
""")
w("app/web/templates/proposal_edit.html", """
{% extends "layout.html" %}
{% block content %}
<div class="toolbar">
  <div class="left">
    <span class="kp-head">КП {{ p.number }} — {{ p.title }}</span>
  </div>
  <div class="right">
    <button onclick="saveKP()">Сохранить</button>
    <a class="btn" target="_blank" href="/proposals/{{ p.id }}/print">Печать</a>
    <a class="btn" target="_blank" href="/proposals/{{ p.id }}/pdf">В PDF</a>
    <a class="btn" href="/contract/editor?customer_id={{ p.customer_id }}&kp_id={{ p.id }}">Редактор договора</a>
  </div>
</div>

<div class="card">
  <h3>Шапка КП</h3>
  <div class="grid2">
    <label>Название:</label><input id="title" value="{{ p.title }}">
    <label>Адрес объекта:</label><input id="site_address" value="{{ p.site_address or p.customer.address_object or '' }}" placeholder="Адрес объекта">
    <label>Бренд/шапка (HTML):</label><textarea id="header_html" rows="3" placeholder="<b>VENT•PRO HVAC</b>">{{ p.header_html or "<b>VENT•PRO HVAC</b>" }}</textarea>
  </div>
</div>

<div class="card">
  <h3>Разделы и позиции (Excel-стиль)</h3>
  <form onsubmit="return addSection()">
    <input id="sec-title" placeholder="Название раздела">
    <button>Добавить раздел</button>
  </form>
  <div id="sections">
    {% for s in sections %}
      <div class="section" data-id="{{ s.id }}">
        <h4>{{ s.position }}. {{ s.title }}</h4>
        <form onsubmit="return addItem({{ s.id }})">
          <input name="name" placeholder="Наименование и тех. характеристика" style="width:26%">
          <input name="note" placeholder="Тип, марка, обозначение" style="width:22%">
          <input name="unit" value="шт" style="width:60px">
          <input name="qty" value="1" style="width:80px">
          <input name="price" value="0" style="width:120px" title="Цена материалов за ед.">
          <input name="price_labor" value="0" style="width:140px" title="Цена монтажа за ед.">
          <button>Добавить позицию</button>
        </form>
        <table class="items" id="items-{{ s.id }}">
          <tr>
            <th>Поз.</th><th>Наименование</th><th>Тип, марка</th><th>Ед.</th><th>Кол-во</th>
            <th>Цена (мат.)</th><th>Мат. всего</th><th>Цена (монтаж)</th><th>Монтаж всего</th><th>Итого</th>
          </tr>
          {% for i in items if i.section_id == s.id %}
            {% set m_total = (i.qty or 0)*(i.price or 0) %}
            {% set l_total = (i.qty or 0)*(i.price_labor or 0) %}
            <tr>
              <td>{{ i.position or "" }}</td>
              <td>{{ i.name }}</td>
              <td>{{ i.note or "" }}</td>
              <td>{{ i.unit }}</td>
              <td>{{ i.qty }}</td>
              <td>{{ "%.0f" % (i.price or 0) }}</td>
              <td>{{ "%.0f" % m_total }}</td>
              <td>{{ "%.0f" % (i.price_labor or 0) }}</td>
              <td>{{ "%.0f" % l_total }}</td>
              <td>{{ "%.0f" % (m_total + l_total) }}</td>
            </tr>
          {% endfor %}
        </table>
      </div>
    {% endfor %}
  </div>
</div>

<script>
async function saveKP(){
  const payload = {
    title: document.getElementById('title').value,
    site_address: document.getElementById('site_address').value,
    header_html: document.getElementById('header_html').value,
  };
  await fetch('/proposals/{{ p.id }}/save', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(payload)});
  alert('Сохранено');
}
async function addSection(){
  const title = document.getElementById('sec-title').value.trim();
  if(!title) return false;
  const r = await fetch('/api/v1/proposals/sections', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({proposal_id: {{ p.id }}, title})});
  if(r.ok) location.reload(); else alert('Ошибка');
  return false;
}
async function addItem(secId){
  const form = event.target;
  const payload = {
    proposal_id: {{ p.id }},
    section_id: secId,
    name: form.querySelector('[name=name]').value,
    note: form.querySelector('[name=note]').value,
    unit: form.querySelector('[name=unit]').value,
    qty: parseFloat(form.querySelector('[name=qty]').value||'1'),
    price: parseFloat(form.querySelector('[name=price]').value||'0'),
    price_labor: parseFloat(form.querySelector('[name=price_labor]').value||'0')
  };
  const r = await fetch('/api/v1/proposals/items', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(payload)});
  if(r.ok) location.reload(); else alert('Ошибка');
  return false;
}
</script>
{% endblock %}
""")
w("app/web/templates/proposal_print.html", """
<!doctype html>
<html><head><meta charset="utf-8"><title>{{ p.number }} — КП</title>
<style>
body{font-family:DejaVu Sans,Arial; margin:20px; color:#111}
.header{display:flex;align-items:center;justify-content:space-between;border-bottom:2px solid #333;padding-bottom:8px;margin-bottom:16px}
.brand{font-weight:bold;font-size:18px}
.kpno{font-size:14px;text-align:right}
.addr{color:#444}
h2{margin:6px 0 12px}
table{border-collapse:collapse;width:100%;font-size:12px}
th,td{border:1px solid #666;padding:6px;vertical-align:top}
.section{margin-top:14px}
.total{font-weight:bold;background:#f5f5f5}
.stage{border:1px solid #666;padding:8px;margin-top:12px}
.small{font-size:11px;color:#444}
</style>
</head>
<body>
<div class="header">
  <div class="brand">{{ p.header_html | safe }}</div>
  <div class="kpno">
    КП № {{ p.number }}<br>
    Клиент: {{ p.customer.fullname }}<br>
    Контакты: {{ p.customer.phone or "" }} {{ p.customer.email or "" }}
  </div>
</div>
<div class="addr"><b>Объект:</b> {{ p.site_address or p.customer.address_object or '' }}</div>

<h2>{{ p.title }}</h2>

{% set equip_total = 0 %}
{% set labor_total = 0 %}

{% for s in sections %}
  <div class="section">
    <h3>Раздел № {{ s.position }}. {{ s.title }}</h3>
    <table>
      <tr>
        <th>Поз.</th><th>Наименование и тех. характеристика</th>
        <th>Тип, марка, обозначение</th><th>Ед.</th><th>Кол-во</th>
        <th>Стоимость оборудования и материалов (на ед.)</th>
        <th>Оборудование и материалы (всего)</th>
        <th>Стоимость монтажных работ (на ед.)</th>
        <th>Монтажные работы (всего)</th>
        <th>Итого в рублях</th>
      </tr>
      {% for i in items if i.section_id == s.id %}
        {% set m_total = (i.qty or 0)*(i.price or 0) %}
        {% set l_total = (i.qty or 0)*(i.price_labor or 0) %}
        {% set row_total = m_total + l_total %}
        {% set equip_total = equip_total + m_total %}
        {% set labor_total = labor_total + l_total %}
        <tr>
          <td>{{ i.position or "" }}</td>
          <td>{{ i.name }}</td>
          <td>{{ i.note or "" }}</td>
          <td>{{ i.unit }}</td>
          <td>{{ i.qty }}</td>
          <td>{{ "%.0f" % (i.price or 0) }}</td>
          <td>{{ "%.0f" % m_total }}</td>
          <td>{{ "%.0f" % (i.price_labor or 0) }}</td>
          <td>{{ "%.0f" % l_total }}</td>
          <td>{{ "%.0f" % row_total }}</td>
        </tr>
      {% endfor %}
    </table>
  </div>
{% endfor %}

{% set grand = equip_total + labor_total %}

<table style="margin-top:10px">
  <tr class="total">
    <td colspan="6">Всего по предложению:</td>
    <td>{{ "%.0f" % equip_total }}</td>
    <td></td>
    <td>{{ "%.0f" % labor_total }}</td>
    <td>{{ "%.0f" % grand }}</td>
  </tr>
</table>

<div class="stage">
  <b>Этапы оплаты:</b>
  <div>1) 100% оборудования и материалов: <b>{{ "%.0f" % equip_total }}</b> руб.</div>
  <div>2) 70% монтажных работ: <b>{{ "%.0f" % (labor_total*0.7) }}</b> руб.</div>
  <div>3) 30% монтажных работ: <b>{{ "%.0f" % (labor_total*0.3) }}</b> руб.</div>
  <div class="small">* Срок действия предложения: 14 календарных дней. Сроки поставки/монтажа и гарантия указываются в договоре.</div>
</div>

</body></html>
""")
w("app/web/templates/proposal_new.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Быстрое создание КП</h2>
<div class="card">
  <form method="post" action="/proposals/new">
    <h3>Выбор клиента</h3>
    <p>Если клиента ещё нет — заполните форму ниже, и он будет создан автоматически.</p>
    <div class="grid2">
      <label>ID существующего клиента:</label><input name="customer_id" placeholder="(опционально)">
      <label>ФИО:</label><input name="fullname" placeholder="Напр.: Иванов Иван Иванович">
      <label>Телефон:</label><input name="phone" placeholder="+7 ...">
      <label>Email:</label><input name="email" placeholder="mail@example.com">
      <label>Адрес регистрации:</label><input name="address" placeholder="Город, улица ...">
      <label>Адрес объекта:</label><input name="address_object" placeholder="Город, улица ...">
    </div>
    <h3>Параметры КП</h3>
    <div class="grid2">
      <label>Название КП:</label><input name="title" value="Коммерческое предложение">
      <label>Адрес объекта (если нужно переопределить):</label><input name="site_address">
      <label>Шапка (HTML):</label><textarea name="header_html" rows="3"><b>VENT•PRO HVAC</b></textarea>
    </div>
    <button>Создать КП</button>
  </form>
</div>
{% endblock %}
""")
w("app/web/templates/templates.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Шаблоны документов (DOCX)</h2>
<p>Для загрузки требуется HTTP Basic (admin/admin по умолчанию).</p>
<form class="card" method="post" action="/templates/upload" enctype="multipart/form-data">
  <input name="name" placeholder="Имя файла, например: Договор.docx" required>
  <input type="file" name="file" required>
  <button>Загрузить</button>
</form>
{% endblock %}
""")
w("app/web/templates/materials.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Материалы</h2>
<form class="card" method="post" action="/materials/create">
  <input name="name" placeholder="Наименование" required>
  <input name="model" placeholder="Модель">
  <input name="unit" placeholder="Ед." value="шт">
  <input name="price_material" placeholder="Цена (мат.)" value="0">
  <input name="price_labor" placeholder="Цена (монтаж)" value="0">
  <button>Добавить</button>
</form>
<form class="card" method="post" action="/materials/import" enctype="multipart/form-data">
  <label>Импорт CSV:</label>
  <input type="file" name="file" required>
  <button>Импортировать</button>
</form>
<form class="card" method="get" action="/materials">
  <input name="q" placeholder="Поиск..." value="{{ q or '' }}">
  <button>Искать</button>
</form>
{% endblock %}
""")
w("app/web/templates/documents.html", """
{% extends "layout.html" %}
{% block content %}
<h2>Документы</h2>
<table>
  <tr><th>Дата</th><th>Тип</th><th>Название</th><th>Номер</th><th>Скачать</th></tr>
  {% for d in docs %}
  <tr>
    <td>{{ d.created_at.strftime('%Y-%m-%d %H:%M') if d.created_at else '' }}</td>
    <td>{{ d.kind }}</td>
    <td>{{ d.title }}</td>
    <td>{{ d.number or '' }}</td>
    <td><a target="_blank" href="/files?path={{ d.filepath | replace('storage\\\\docs\\\\','') | replace('storage/docs/','') }}">{{ d.filepath.split('/')[-1].split('\\\\')[-1] }}</a></td>
  </tr>
  {% endfor %}
</table>
{% endblock %}
""")
w("app/web/templates/contract_editor.html", """
{% extends "layout.html" %}
{% block content %}
<div class="toolbar">
  <div class="left"><b>Редактор договора</b></div>
  <div class="right">
    <button onclick="saveCtx()">Сохранить параметры</button>
    <a class="btn" target="_blank" id="btn-docx">Сформировать DOCX</a>
  </div>
</div>
<div class="grid2 card">
  <label>Дата договора:</label><input id="CONTRACT_DATE" value="{{ ctx.CONTRACT_DATE }}">
  <label>Срок поставки (дн.):</label><input id="DELIVERY_DAYS" value="{{ ctx.DELIVERY_DAYS or '' }}">
  <label>Срок монтажа (дн.):</label><input id="INSTALL_DAYS" value="{{ ctx.INSTALL_DAYS or '' }}">
  <label>Гарантия (мес.):</label><input id="WARRANTY_MONTHS" value="{{ ctx.WARRANTY_MONTHS or '' }}">
  <label>Срок действия КП (дн.):</label><input id="KP_VALID_DAYS" value="{{ ctx.KP_VALID_DAYS or '' }}">
  <label>Предоплата (%):</label><input id="PREPAYMENT_PCT" value="{{ ctx.PREPAYMENT_PCT or '' }}">
</div>
<iframe id="preview" style="width:100%;height:70vh;border:1px solid #334155"></iframe>
<script>
const params = new URLSearchParams(window.location.search);
const customer_id = params.get('customer_id');
const estimate_id = params.get('estimate_id');
const kp_id = params.get('kp_id');

function reloadPreview(){
  const q = new URLSearchParams({
    customer_id, estimate_id, kp_id,
    CONTRACT_DATE: document.getElementById('CONTRACT_DATE').value,
    DELIVERY_DAYS: document.getElementById('DELIVERY_DAYS').value,
    INSTALL_DAYS: document.getElementById('INSTALL_DAYS').value,
    WARRANTY_MONTHS: document.getElementById('WARRANTY_MONTHS').value,
    KP_VALID_DAYS: document.getElementById('KP_VALID_DAYS').value,
    PREPAYMENT_PCT: document.getElementById('PREPAYMENT_PCT').value,
  });
  document.getElementById('preview').src = '/print/contract?'+q.toString();
  document.getElementById('btn-docx').href = '/contract/docx?'+q.toString();
}
reloadPreview();

async function saveCtx(){
  const overrides = {
    CONTRACT_DATE: document.getElementById('CONTRACT_DATE').value,
    DELIVERY_DAYS: document.getElementById('DELIVERY_DAYS').value,
    INSTALL_DAYS: document.getElementById('INSTALL_DAYS').value,
    WARRANTY_MONTHS: document.getElementById('WARRANTY_MONTHS').value,
    KP_VALID_DAYS: document.getElementById('KP_VALID_DAYS').value,
    PREPAYMENT_PCT: document.getElementById('PREPAYMENT_PCT').value,
  };
  await fetch('/api/v1/contracts/save-context?customer_id='+customer_id+'&estimate_id='+ (estimate_id||'') +'&kp_id='+ (kp_id||''), {
    method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({overrides})
  });
  alert('Сохранено'); reloadPreview();
}
</script>
{% endblock %}
""")
w("app/web/templates/print_contract.html", """
<!doctype html>
<html><head><meta charset="utf-8"><title>Договор</title>
<style>
body{font-family:DejaVu Sans,Arial} h1{font-size:18px} .row{margin:6px 0}
.header{display:flex;justify-content:space-between;align-items:center;border-bottom:2px solid #333;padding-bottom:6px;margin-bottom:8px}
.brand{font-weight:bold}
</style></head>
<body>
<div class="header">
  <div class="brand">VENT•PRO — HVAC Solutions</div>
  <div>Договор № {{ ctx.CONTRACT_NO }} от {{ ctx.CONTRACT_DATE }}</div>
</div>
<div class="row"><b>Заказчик:</b> {{ ctx.CUSTOMER_FIO }}</div>
<div class="row"><b>Исполнитель:</b> {{ ctx.EXECUTOR_FIO }}</div>
<div class="row"><b>Объект:</b> {{ ctx.OBJECT_ADDR }}</div>
<hr>
<p>Стоимость оборудования и материалов: <b>{{ ctx.SUM_EQUIP }}</b> руб.</p>
<p>Стоимость монтажных работ: <b>{{ ctx.SUM_WORK }}</b> руб.</p>
<p>Итого: <b>{{ ctx.SUM_TOTAL }}</b> руб.</p>
<p>Предоплата: <b>{{ ctx.PREPAYMENT_PCT or "по договоренности" }}</b>%</p>
<p>Срок поставки: <b>{{ ctx.DELIVERY_DAYS or "-" }}</b> дней, срок монтажа: <b>{{ ctx.INSTALL_DAYS or "-" }}</b> дней, гарантия: <b>{{ ctx.WARRANTY_MONTHS or "-" }}</b> мес.</p>
<ol>
  <li>Этап 1: {{ ctx.STAGE1 }} руб.</li>
  <li>Этап 2: {{ ctx.STAGE2 }} руб.</li>
  <li>Этап 3: {{ ctx.STAGE3 }} руб.</li>
</ol>
<p>Срок действия КП: {{ ctx.KP_VALID_DAYS or "-" }} дней.</p>
</body></html>
""")

# ---------- main.py ----------
w("app/main.py", """
from fastapi import FastAPI, Request, UploadFile, File, Form, Depends, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from sqlalchemy import select
from xhtml2pdf import pisa
from io import BytesIO

from app.db.init_db import init_db
from app.db.session import SessionLocal
from app.models import Customer, Estimate, EstimateSection, EstimateItem, Material, TemplateRec, DocumentRec, Owner, Proposal, ProposalSection, ProposalItem
from app.api.routers.customers import router as customers_router
from app.api.routers.estimates import router as estimates_router
from app.api.routers.materials import router as materials_router
from app.api.routers.categories import router as categories_router
from app.api.routers.contracts import router as contracts_router
from app.api.routers.templates import router as tpl_router
from app.api.routers.files import router as files_router
from app.api.routers.aihelper import router as ai_router
from app.api.routers.proposals import router as proposals_api
from app.config import STATIC_DIR, HTML_TPL_DIR, DOCS_DIR, TEMPLATES_DIR
from app.services.numbering import generate_proposal_number
from app.services.payments import estimate_totals, payment_schedule_100_70_30, proposal_totals
from app.services.docx_merge import render_docx

app = FastAPI(title="KP Suite ULTRA")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
templates = Jinja2Templates(directory=HTML_TPL_DIR)

@app.on_event("startup")
def on_start(): init_db()

# API
app.include_router(files_router)
app.include_router(customers_router)
app.include_router(estimates_router)
app.include_router(materials_router)
app.include_router(categories_router)
app.include_router(contracts_router)
app.include_router(tpl_router)
app.include_router(ai_router)
app.include_router(proposals_api)

# UI
@app.get("/", response_class=HTMLResponse)
def index(request: Request): return templates.TemplateResponse("index.html", {"request": request})

@app.get("/customers", response_class=HTMLResponse)
def page_customers(request: Request):
    db: Session = SessionLocal()
    customers = db.execute(select(Customer)).scalars().all()
    db.close()
    return templates.TemplateResponse("customers.html", {"request": request, "customers": customers})

# ---------- CUSTOMER: create/save/full card ----------
@app.get("/customers/create", response_class=HTMLResponse)
def customers_create_get(request: Request):
    return RedirectResponse(url="/customers", status_code=303)

@app.post("/customers/create")
async def customers_create(fullname: str = Form(...), phone: str = Form(None), email: str = Form(None),
                           address: str = Form(None), address_object: str = Form(None), passport: str = Form(None)):
    db = SessionLocal()
    c = Customer(fullname=fullname, phone=phone, email=email, address=address, address_object=address_object, passport=passport)
    db.add(c); db.commit(); db.close()
    return RedirectResponse(url="/customers", status_code=303)

@app.get("/customers/card/{cid}", response_class=HTMLResponse)
def page_customer_card_full(request: Request, cid: int):
    db = SessionLocal()
    c = db.get(Customer, cid)
    if not c:
        db.close()
        return RedirectResponse(url="/customers", status_code=303)
    latest_kp = db.execute(select(Proposal).where(Proposal.customer_id==cid).order_by(Proposal.created_at.desc())).scalars().first()
    props = db.execute(select(Proposal).where(Proposal.customer_id==cid).order_by(Proposal.created_at.desc())).scalars().all()
    db.close()
    return templates.TemplateResponse("customer_card_full.html", {"request": request, "c": c, "latest_kp": latest_kp, "proposals": props})

@app.post("/customers/{cid}/save")
async def customers_save(cid: int, request: Request):
    form = await request.form()
    db = SessionLocal()
    c = db.get(Customer, cid)
    if not c:
        db.close()
        return RedirectResponse(url="/customers", status_code=303)
    c.fullname = form.get("fullname") or c.fullname
    c.passport = form.get("passport") or c.passport
    c.address = form.get("address") or c.address
    c.address_object = form.get("address_object") or c.address_object
    c.phone = form.get("phone") or c.phone
    c.email = form.get("email") or c.email
    db.add(c); db.commit(); db.close()
    return RedirectResponse(url=f"/customers/card/{cid}", status_code=303)

@app.post("/customers/{cid}/make-contract-latest-kp")
def make_contract_latest_kp(cid: int):
    db = SessionLocal()
    kp = db.execute(select(Proposal).where(Proposal.customer_id==cid).order_by(Proposal.created_at.desc())).scalars().first()
    db.close()
    if kp:
        return RedirectResponse(url=f"/contract/editor?customer_id={cid}&kp_id={kp.id}", status_code=303)
    else:
        return RedirectResponse(url=f"/contract/editor?customer_id={cid}", status_code=303)

# ---------- PROPOSALS (КП) UI ----------
@app.get("/proposals", response_class=HTMLResponse)
def list_proposals(request: Request):
    db: Session = SessionLocal()
    rows = db.execute(select(Proposal).order_by(Proposal.created_at.desc())).scalars().all()
    for r in rows: _ = r.customer
    db.close()
    return templates.TemplateResponse("proposals.html", {"request": request, "proposals": rows})

@app.get("/proposals/new", response_class=HTMLResponse)
def new_proposal_page(request: Request):
    return templates.TemplateResponse("proposal_new.html", {"request": request})

@app.post("/proposals/new")
async def new_proposal_create(request: Request):
    form = await request.form()
    customer_id = form.get("customer_id")
    db: Session = SessionLocal()
    from app.models import Customer, Proposal
    number = generate_proposal_number(db)
    # клиент
    c = None
    if customer_id:
        try: c = db.get(Customer, int(customer_id))
        except: c = None
    if not c:
        c = Customer(
            fullname=form.get("fullname") or "Новый клиент",
            phone=form.get("phone") or None,
            email=form.get("email") or None,
            address=form.get("address") or None,
            address_object=form.get("address_object") or None,
        )
        db.add(c); db.commit(); db.refresh(c)
    p = Proposal(
        customer_id=c.id, number=number,
        title=form.get("title") or "Коммерческое предложение",
        site_address=form.get("site_address") or c.address_object,
        header_html=form.get("header_html") or "<b>VENT•PRO HVAC</b>",
    )
    db.add(p); db.commit(); db.refresh(p); db.close()
    return RedirectResponse(url=f"/proposals/{p.id}/edit", status_code=303)

@app.get("/proposals/create")
def create_proposal_page(customer_id: int):
    db: Session = SessionLocal()
    number = generate_proposal_number(db)
    p = Proposal(customer_id=customer_id, number=number, title="Коммерческое предложение")
    db.add(p); db.commit(); db.refresh(p); db.close()
    return RedirectResponse(url=f"/proposals/{p.id}/edit", status_code=303)

@app.get("/proposals/{pid}/edit", response_class=HTMLResponse)
def edit_proposal(request: Request, pid: int):
    db: Session = SessionLocal()
    p = db.get(Proposal, pid)
    sections = db.execute(select(ProposalSection).where(ProposalSection.proposal_id==pid).order_by(ProposalSection.position)).scalars().all()
    items = db.execute(select(ProposalItem).where(ProposalItem.proposal_id==pid)).scalars().all()
    db.close()
    return templates.TemplateResponse("proposal_edit.html", {"request": request, "p": p, "sections": sections, "items": items})

@app.post("/proposals/{pid}/save")
async def save_proposal(pid: int, request: Request):
    data = await request.json()
    db: Session = SessionLocal()
    p = db.get(Proposal, pid)
    p.title = data.get("title", p.title)
    p.site_address = data.get("site_address", p.site_address)
    p.header_html = data.get("header_html", p.header_html)
    db.add(p); db.commit(); db.close()
    return {"status":"ok"}

@app.get("/proposals/{pid}/print", response_class=HTMLResponse)
def print_proposal(request: Request, pid: int):
    db: Session = SessionLocal()
    p = db.get(Proposal, pid)
    sections = db.execute(select(ProposalSection).where(ProposalSection.proposal_id==pid).order_by(ProposalSection.position)).scalars().all()
    items = db.execute(select(ProposalItem).where(ProposalItem.proposal_id==pid)).scalars().all()
    db.close()
    return templates.TemplateResponse("proposal_print.html", {"request": request, "p": p, "sections": sections, "items": items})

@app.get("/proposals/{pid}/pdf")
def pdf_proposal(pid: int):
    req = Request(scope={"type": "http"})
    html_resp = print_proposal(req, pid=pid)
    from starlette.templating import _TemplateResponse
    if isinstance(html_resp, _TemplateResponse):
        html_resp.render(); html = html_resp.body.decode("utf-8")
    else: html = str(html_resp)
    out = DOCS_DIR / f"KP_{pid}.pdf"
    buf = BytesIO(); pisa.CreatePDF(src=html, dest=buf); out.write_bytes(buf.getvalue())
    return FileResponse(str(out), filename=out.name, media_type="application/pdf")

# ---------- CONTRACT EDITOR & PRINT ----------
@app.get("/contract/editor", response_class=HTMLResponse)
def contract_editor(request: Request, customer_id: int, estimate_id: int | None = None, kp_id: int | None = None):
    from datetime import date
    ctx = {"CONTRACT_DATE": date.today().strftime("%d.%m.%Y")}
    return templates.TemplateResponse("contract_editor.html", {"request": request, "ctx": ctx})

@app.get("/print/contract", response_class=HTMLResponse)
def print_contract(request: Request, customer_id: int, estimate_id: int | None = None, kp_id: int | None = None,
                   CONTRACT_DATE: str | None = None, DELIVERY_DAYS: str | None = None, INSTALL_DAYS: str | None = None,
                   WARRANTY_MONTHS: str | None = None, KP_VALID_DAYS: str | None = None, PREPAYMENT_PCT: str | None = None):
    db: Session = SessionLocal()
    cust = db.get(Customer, customer_id)
    owner = db.get(Owner, 1) or Owner(id=1, display_name="Исполнитель"); db.add(owner); db.commit(); db.refresh(owner)
    est = db.get(Estimate, estimate_id) if estimate_id else None
    kp = db.get(Proposal, kp_id) if kp_id else None
    from app.services.numbering import generate_contract_number
    from app.services.payments import estimate_totals, proposal_totals
    totals = {"equipment":0,"labor":0,"total":0}
    if est: totals = estimate_totals(est)
    elif kp: totals = proposal_totals(kp)
    ctx = {
        "CONTRACT_NO": generate_contract_number(db),
        "CONTRACT_DATE": CONTRACT_DATE or "",
        "CUSTOMER_FIO": cust.fullname or "",
        "OBJECT_ADDR": (est.site_address if est else (kp.site_address if kp else "")) or cust.address_object or "",
        "EXECUTOR_FIO": owner.fullname or owner.display_name or "Исполнитель",
        "SUM_EQUIP": f"{totals['equipment']:,.0f}".replace(",", " "),
        "SUM_WORK": f"{totals['labor']:,.0f}".replace(",", " "),
        "SUM_TOTAL": f"{totals['total']:,.0f}".replace(",", " "),
        "STAGE1": "", "STAGE2": "", "STAGE3": "",
        "DELIVERY_DAYS": DELIVERY_DAYS, "INSTALL_DAYS": INSTALL_DAYS, "WARRANTY_MONTHS": WARRANTY_MONTHS,
        "KP_VALID_DAYS": KP_VALID_DAYS, "PREPAYMENT_PCT": PREPAYMENT_PCT,
    }
    db.close()
    return templates.TemplateResponse("print_contract.html", {"request": request, "ctx": ctx})

@app.get("/contract/docx")
def contract_docx(customer_id: int, estimate_id: int | None = None, kp_id: int | None = None, **kwargs):
    db: Session = SessionLocal()
    from app.api.routers.contracts import build_ctx
    number, ctx, _, _ = build_ctx(db, customer_id, estimate_id, kp_id, kwargs)
    tpl = TEMPLATES_DIR / "Договор.docx"
    if not tpl.is_file():
        raise HTTPException(400, "Нет шаблона Договор.docx")
    out = DOCS_DIR / f"Договор_{number}.docx"
    render_docx(str(tpl), str(out), ctx)
    return FileResponse(str(out), filename=out.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---------- MATERIALS UI ----------
@app.get("/materials", response_class=HTMLResponse)
def materials_page(request: Request):
    return templates.TemplateResponse("materials.html", {"request": request})

# ---------- TEMPLATES UI ----------
@app.get("/templates", response_class=HTMLResponse)
def page_templates(request: Request):
    return templates.TemplateResponse("templates.html", {"request": request})

# ---------- Documents UI ----------
@app.get("/documents", response_class=HTMLResponse)
def page_documents(request: Request):
    db: Session = SessionLocal()
    docs = db.execute(select(DocumentRec).order_by(DocumentRec.created_at.desc())).scalars().all()
    db.close()
    return templates.TemplateResponse("documents.html", {"request": request, "docs": docs})

@app.get("/health")
def health(): return {"status": "ok"}
""")

# ---------- helpers: seed_catalog + make_templates ----------
w("seed_catalog.py", r"""
# Заполняет категории и позиции (сантехника, электрика, оборудование, изоляция, работы).
from app.db.session import SessionLocal
from sqlalchemy import select
from app.models import Category, Material

db = SessionLocal()

cats = [
    ("Сантехника", None, "materials"),
    ("Электрика", None, "materials"),
    ("Оборудование", None, "materials"),
    ("Изоляция", None, "materials"),
    ("Работы", None, "works"),
]
existing = {c.name: c for c in db.execute(select(Category)).scalars().all()}
for name, parent_id, kind in cats:
    if name not in existing:
        db.add(Category(name=name, parent_id=parent_id, kind=kind))
db.commit()
cat_map = {c.name: c.id for c in db.execute(select(Category)).scalars().all()}

materials = [
    # Сантехника
    dict(name="Смеситель для раковины", model="универсальный", unit="шт", price_material=4500, price_labor=1200, category_id=cat_map["Сантехника"]),
    dict(name="Унитаз-компакт", unit="шт", price_material=14500, price_labor=3500, category_id=cat_map["Сантехника"]),
    dict(name="Раковина керамическая 60 см", unit="шт", price_material=6200, price_labor=2500, category_id=cat_map["Сантехника"]),
    dict(name="Сифон ПВХ 50", unit="шт", price_material=350, price_labor=300, category_id=cat_map["Сантехника"]),
    dict(name="Гибкая подводка 1/2\" 60 см", unit="шт", price_material=250, price_labor=0, category_id=cat_map["Сантехника"]),
    # Электрика
    dict(name="Кабель ВВГнг-LS 3x2.5", unit="м", price_material=120, price_labor=50, category_id=cat_map["Электрика"]),
    dict(name="Автоматический выключатель 16A", unit="шт", price_material=480, price_labor=200, category_id=cat_map["Электрика"]),
    dict(name="Розетка двойная", unit="шт", price_material=350, price_labor=250, category_id=cat_map["Электрика"]),
    dict(name="Щит распределительный 12 мод.", unit="шт", price_material=1900, price_labor=1200, category_id=cat_map["Электрика"]),
    dict(name="Стабилизатор 5 кВт", unit="шт", price_material=9800, price_labor=1500, category_id=cat_map["Электрика"]),
    # Оборудование (HVAC)
    dict(name="Вентилятор канальный 315 мм", unit="шт", price_material=12500, price_labor=3500, category_id=cat_map["Оборудование"]),
    dict(name="Воздуховод оцинкованный Ø250", unit="м", price_material=800, price_labor=200, category_id=cat_map["Оборудование"]),
    dict(name="Сплит-система 3,5 кВт", model="LG", unit="шт", price_material=52000, price_labor=15000, category_id=cat_map["Оборудование"]),
    dict(name="Фильтр кассетный G4", unit="шт", price_material=2100, price_labor=400, category_id=cat_map["Оборудование"]),
    dict(name="Решетка вентиляционная 200x200", unit="шт", price_material=650, price_labor=150, category_id=cat_map["Оборудование"]),
    # Изоляция
    dict(name="Теплоизоляция рулонная 50 мм", unit="м2", price_material=280, price_labor=180, category_id=cat_map["Изоляция"]),
    dict(name="Теплоизоляция труб 25 мм", unit="м", price_material=90, price_labor=60, category_id=cat_map["Изоляция"]),
    dict(name="Изоляция воздуховодов фольгированная", unit="м2", price_material=420, price_labor=220, category_id=cat_map["Изоляция"]),
    # Работы (только труд)
    dict(name="Демонтаж вентиляционного оборудования", unit="час", price_material=0, price_labor=1200, category_id=cat_map["Работы"]),
    dict(name="Демонтаж труб/воздуховодов", unit="м", price_material=0, price_labor=300, category_id=cat_map["Работы"]),
    dict(name="Разгрузочно-погрузочные работы", unit="час", price_material=0, price_labor=900, category_id=cat_map["Работы"]),
    dict(name="Подъём на этаж (без лифта)", unit="этаж", price_material=0, price_labor=400, category_id=cat_map["Работы"]),
    dict(name="Уборочные работы после монтажа", unit="час", price_material=0, price_labor=600, category_id=cat_map["Работы"]),
]

exist = {(m.name or "", m.model or "") for m in db.execute(select(Material)).scalars().all()}
ins = 0
for m in materials:
    key = (m["name"] or "", m.get("model") or "")
    if key in exist: continue
    db.add(Material(**m)); ins += 1
db.commit(); db.close()
print(f"Catalog seeded: +{ins} items")
""")

w("make_templates.py", r"""
# Создаёт шаблон договора с плейсхолдерами и пример КП (DOCX) в проекте.
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

BASE = Path(__file__).resolve().parent
TPL_DIR = BASE / "storage" / "templates" / "contracts"
DOCS_DIR = BASE / "storage" / "docs"
TPL_DIR.mkdir(parents=True, exist_ok=True)
DOCS_DIR.mkdir(parents=True, exist_ok=True)

def heading(doc, text, size=16, bold=True, align="center"):
    p = doc.add_paragraph()
    run = p.add_run(text); run.bold = bold; run.font.size = Pt(size)
    p.alignment = {"center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                   "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                   "right": WD_PARAGRAPH_ALIGNMENT.RIGHT}[align]

def kv(doc, label, value):
    p = doc.add_paragraph(); r1 = p.add_run(label + " "); r1.bold = True; p.add_run(value)

# ---------- Договор ----------
dog = Document()
heading(dog, "ДОГОВОР № [[CONTRACT_NO]]", 18, True, "center")
heading(dog, "от [[CONTRACT_DATE]] г.", 12, False, "center")
dog.add_paragraph()
kv(dog, "Заказчик:", "[[CUSTOMER_FIO]]")
kv(dog, "Паспорт:", "[[CUSTOMER_PASSPORT]]")
kv(dog, "Адрес регистрации:", "[[CUSTOMER_ADDR_REG]]")
kv(dog, "Телефон:", "[[CUSTOMER_PHONE]]")
kv(dog, "Email:", "[[CUSTOMER_EMAIL]]")
dog.add_paragraph()
kv(dog, "Исполнитель:", "[[EXECUTOR_FIO]]")
kv(dog, "Паспорт/реквизиты:", "[[EXECUTOR_PASSPORT]]")
kv(dog, "Адрес:", "[[EXECUTOR_ADDR]]")
kv(dog, "Телефон:", "[[EXECUTOR_PHONE]]")
kv(dog, "Email:", "[[EXECUTOR_EMAIL]]")
kv(dog, "ИНН:", "[[EXECUTOR_INN]]")
kv(dog, "ОГРН:", "[[EXECUTOR_OGRN]]")
dog.add_paragraph(); kv(dog, "Объект:", "[[OBJECT_ADDR]]")

dog.add_paragraph()
heading(dog, "1. Стоимость работ и материалов", 14, True, "left")
t = dog.add_table(rows=4, cols=2)
t.rows[0].cells[0].text = "Позиция";             t.rows[0].cells[1].text = "Сумма, руб."
t.rows[1].cells[0].text = "Оборудование и материалы"; t.rows[1].cells[1].text = "[[SUM_EQUIP]]"
t.rows[2].cells[0].text = "Монтажные работы";    t.rows[2].cells[1].text = "[[SUM_WORK]]"
t.rows[3].cells[0].text = "ИТОГО";               t.rows[3].cells[1].text = "[[SUM_TOTAL]]"

dog.add_paragraph()
heading(dog, "2. Порядок расчётов", 14, True, "left")
dog.add_paragraph("2.1. Предоплата составляет [[PREPAYMENT_PCT]]% от общей стоимости.")
pt = dog.add_table(rows=4, cols=2)
pt.rows[0].cells[0].text = "Этап";   pt.rows[0].cells[1].text = "Сумма, руб."
pt.rows[1].cells[0].text = "Этап 1"; pt.rows[1].cells[1].text = "[[STAGE1]]"
pt.rows[2].cells[0].text = "Этап 2"; pt.rows[2].cells[1].text = "[[STAGE2]]"
pt.rows[3].cells[0].text = "Этап 3"; pt.rows[3].cells[1].text = "[[STAGE3]]"

dog.add_paragraph()
heading(dog, "3. Сроки и гарантия", 14, True, "left")
dog.add_paragraph("3.1. Срок поставки: [[DELIVERY_DAYS]] календарных дней.")
dog.add_paragraph("3.2. Срок монтажа: [[INSTALL_DAYS]] календарных дней.")
dog.add_paragraph("3.3. Гарантия: [[WARRANTY_MONTHS]] месяцев.")
dog.add_paragraph("3.4. Срок действия КП: [[KP_VALID_DAYS]] дней.")

dog.add_paragraph()
heading(dog, "4. Прочие условия", 14, True, "left")
dog.add_paragraph("4.1. Стороны обязуются выполнять условия настоящего договора добросовестно.")
dog.add_paragraph("4.2. Споры разрешаются путём переговоров.")

dog.add_paragraph()
heading(dog, "5. Реквизиты и подписи сторон", 14, True, "left")
st = dog.add_table(rows=2, cols=2)
st.cell(0,0).text = "Заказчик: [[CUSTOMER_FIO]]\\n_________________/________________"
st.cell(0,1).text = "Исполнитель: [[EXECUTOR_FIO]]\\n_________________/________________"
st.cell(1,0).text = "Тел.: [[CUSTOMER_PHONE]]  Email: [[CUSTOMER_EMAIL]]"
st.cell(1,1).text = "Тел.: [[EXECUTOR_PHONE]]  Email: [[EXECUTOR_EMAIL]]"

dog_out = TPL_DIR / "Договор.docx"
dog.save(dog_out)

# ---------- Пример КП (вне системы, просто образец) ----------
kp = Document()
def add_kv(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(label + " "); r1.bold = True
    p.add_run(value)

heading(kp, "VENT•PRO — HVAC Solutions", 16, True, "left")
heading(kp, "Коммерческое предложение", 18, True, "center")
add_kv(kp, "Клиент:", "ООО «Пример», ИНН 7700000000")
add_kv(kp, "Контакт:", "Иванов Иван Иванович, +7 (900) 000-00-00, example@mail.ru")
add_kv(kp, "Объект:", "г. Москва, ул. Примерная, д. 1, корп. 2")

kp.add_paragraph()
heading(kp, "Состав работ и оборудования", 14, True, "left")

tbl = kp.add_table(rows=1, cols=7)
hdr = tbl.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "Раздел", "Поз.", "Наименование"
hdr[3].text, hdr[4].text, hdr[5].text, hdr[6].text = "Ед.", "Кол-во", "Цена", "Сумма"

def add(razdel, poz, name, unit, qty, price):
    r = tbl.add_row().cells
    r[0].text, r[1].text, r[2].text = razdel, poz, name
    r[3].text, r[4].text = unit, str(qty)
    r[5].text = f"{price:,.0f}".replace(",", " ")
    r[6].text = f"{qty*price:,.0f}".replace(",", " ")

add("1. Вентиляция", "1.1", "Вентилятор осевой 250 мм", "шт", 2, 12500)
add("1. Вентиляция", "1.2", "Канал воздуховод оцинкованный Ø250", "м", 15, 800)
add("2. Кондиционирование", "2.1", "Сплит-система 3,5 кВт (LG)", "шт", 1, 52000)
add("2. Кондиционирование", "2.2", "Монтаж сплит-системы", "компл", 1, 15000)

items = [(2,12500),(15,800),(1,52000),(1,15000)]
total = sum(q*p for q,p in items)
p = kp.add_paragraph(); p.add_run("ИТОГО по КП: ").bold = True; p.add_run(f"{total:,.0f} руб.").bold = True
kp.add_paragraph("Срок действия предложения: 14 календарных дней.")
kp.add_paragraph("Срок поставки: 10 дней. Срок монтажа: 5 дней.")
kp.add_paragraph("Гарантия на оборудование: 24 месяца.")

kp_out = DOCS_DIR / "КП_шаблон_с_примером.docx"
kp.save(kp_out)

print("OK:", dog_out, "и", kp_out)
""")

print("\n== DONE ==")
print("1) Откройте папку kp_suite_ultra и запустите: run_kp.bat")
print("2) После запуска: http://127.0.0.1:8000/")
print("3) Для шаблонов и каталога: .venv\\Scripts\\python.exe make_templates.py  и  .venv\\Scripts\\python.exe seed_catalog.py")
